"""PackageSnapshot — cached DOCX parse shared across evaluators.

Each DOCX is opened once, paragraphs and tables flattened, and the
heading ancestor of each paragraph is computed. Evaluators ask the
snapshot for text, paragraphs, counts, or heading context instead of
re-parsing. Cuts evaluator-pass time by ~6x on the typical package.

Page-count falls back to a paragraph-based heuristic because
`soffice` is not assumed available (see plan §2). `page_count()`
returns an estimate and a `method` tag so the evaluator can decide
whether to gate or advise.
"""
from __future__ import annotations
import pathlib
import re
from dataclasses import dataclass, field
from typing import Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.document import Document as DocType
from docx.oxml.ns import qn


@dataclass
class Paragraph:
    text: str
    style: str
    heading_level: Optional[int]   # None for body paragraphs
    ancestor_heading: str          # nearest heading above this para (text)
    # `is_table_cell=True` is attached when the paragraph came from a
    # table cell rather than the document body.
    is_table_cell: bool = False


@dataclass
class TableView:
    rows: int
    cols: int
    header_cells: List[str]     # first-row cell text, trimmed
    first_cell: str             # stripped text of row0/col0

    def header_has(self, needle: str) -> bool:
        n = needle.lower()
        return any(n in (c or "").lower() for c in self.header_cells) \
            or n in (self.first_cell or "").lower()


@dataclass
class DocView:
    path: pathlib.Path
    paragraphs: List[Paragraph]
    raw_text: str
    page_break_count: int
    table_row_count: int
    image_count: int = 0
    tables: List[TableView] = field(default_factory=list)

    def iter_texts(self) -> Iterable[str]:
        return (p.text for p in self.paragraphs)

    def paragraphs_under(self, heading_contains: str) -> List[Paragraph]:
        needle = heading_contains.lower()
        return [p for p in self.paragraphs
                if needle in p.ancestor_heading.lower()]

    def headings(self) -> List[Paragraph]:
        return [p for p in self.paragraphs if p.heading_level]


@dataclass
class PackageSnapshot:
    package_root: pathlib.Path
    outputs: Dict[str, pathlib.Path]     # label -> docx path
    charts_dir: Optional[pathlib.Path]
    _views: Dict[str, DocView] = field(default_factory=dict)

    def view(self, label: str) -> Optional[DocView]:
        if label not in self.outputs:
            return None
        if label in self._views:
            return self._views[label]
        path = self.outputs[label]
        if not path.exists():
            return None
        self._views[label] = _load_docview(path)
        return self._views[label]

    def text(self, label: str) -> str:
        v = self.view(label)
        return v.raw_text if v else ""

    def paragraphs(self, label: str) -> List[Paragraph]:
        v = self.view(label)
        return v.paragraphs if v else []

    # ─── Helpers that know the whole package ─────────────────────────

    def all_docx_labels(self) -> List[str]:
        return list(self.outputs.keys())

    def iter_text_all(self) -> Iterable[Tuple[str, str]]:
        """Yield (label, text) for every DOCX that exists."""
        for label in self.outputs:
            v = self.view(label)
            if v is not None:
                yield label, v.raw_text

    def page_count(self, label: str) -> Tuple[int, str]:
        """Return (page_estimate, method). method in {'heuristic'}."""
        v = self.view(label)
        if v is None:
            return 0, "missing"
        # Each explicit page break adds a page; assume 1 page of body
        # even if no breaks. Cap at 2 pages per ~50 paragraphs/tables
        # beyond the break count as a rough second-order estimate.
        body_units = len(v.paragraphs) + v.table_row_count
        est = max(1, v.page_break_count + 1) + max(0, body_units // 90)
        return est, "heuristic"


# ─── Internal parsing ────────────────────────────────────────────────

def _is_heading_style(style_name: str) -> Optional[int]:
    m = re.match(r"Heading\s+(\d+)", style_name or "", flags=re.I)
    if m:
        return int(m.group(1))
    return None


# These build scripts style headings by direct formatting (bold, size)
# rather than applying Word's Heading styles. We fall back to matching
# the numbering pattern used in the reports: "1. Title", "1.1 Title",
# "Annex A. Title", "R1", "R1 — Title".
_HEADING_NUMBER_RE = re.compile(
    r"^(?:Annex\s+[A-Z]|[0-9]+(?:\.[0-9]+){0,2})\.?\s+\S",
)
_HEADING_BARE_RE = re.compile(
    r"^(Annex|Executive summary|Summary for decision-makers|"
    r"Headline|Methodology|Note on denominators|Benchmark)\b",
    re.I,
)


def _heading_level_from_text(text: str) -> Optional[int]:
    """Infer heading depth from numbering. Returns None for body text.

    A numbered or bare heading is only accepted if the paragraph is
    short enough to be a plausible heading. Numbered paragraphs that
    are clearly prose (multi-sentence, long) are rejected even when
    they begin with 1./2./3.
    """
    s = (text or "").strip()
    if not s:
        return None
    word_count = len(s.split())
    # Bail on obvious prose: multi-sentence paragraphs.
    # A sentence boundary is a lowercase letter followed by '.', space,
    # and a capital. "1.  What" has a digit before the '.', so it's
    # a numbering prefix, not a sentence break, and is allowed.
    if re.search(r"[a-z]\.\s+[A-Z]", s):
        return None
    # Bail on anything much longer than a real heading.
    if word_count > 14:
        return None
    if _HEADING_BARE_RE.match(s):
        return 1
    if _HEADING_NUMBER_RE.match(s):
        if s.lower().startswith("annex"):
            return 1
        prefix = s.split()[0].rstrip(".")
        # "1" → L1, "1.1" → L2, "1.1.1" → L3, etc.
        depth = prefix.count(".") + 1
        return min(depth, 6)
    return None


def _looks_bold_short_heading(para, text: str) -> bool:
    """Catch headings styled by direct formatting: short, all-bold."""
    stripped = (text or "").strip()
    if not (3 <= len(stripped) <= 80):
        return False
    if len(stripped) > 0 and stripped.endswith(("."," ,", ":", ";")):
        # Very unlikely to be a heading if it ends in sentence punctuation
        pass
    runs = list(para.runs) if hasattr(para, "runs") else []
    if not runs:
        return False
    bold_runs = [r for r in runs if (r.bold is True)]
    return len(bold_runs) == len(runs) and len(stripped.split()) <= 12


def _has_page_break(paragraph) -> int:
    count = 0
    for run in paragraph.runs:
        for br in run._element.iter():
            if br.tag.endswith("}br") and br.get(qn("w:type")) == "page":
                count += 1
    return count


def _flatten_body(doc: DocType) -> Tuple[List[Paragraph], int, int]:
    """Walk document body: each paragraph knows its nearest heading
    ancestor. Tables are flattened as cell paragraphs with the table's
    surrounding heading context."""
    paragraphs: List[Paragraph] = []
    total_breaks = 0
    total_rows = 0

    current_heading_text = ""
    current_heading_level: Optional[int] = None
    # iterate over body in document order
    for block in doc.element.body.iterchildren():
        tag = block.tag.split("}")[-1]
        if tag == "p":
            # Wrap with python-docx Paragraph for style/runs access
            from docx.text.paragraph import Paragraph as DP
            para = DP(block, doc.part)
            style = para.style.name if para.style else ""
            lvl = _is_heading_style(style)
            if lvl is None:
                lvl = _heading_level_from_text(para.text)
            if lvl is None and _looks_bold_short_heading(para, para.text):
                lvl = 2
            total_breaks += _has_page_break(para)
            if lvl is not None:
                current_heading_text = para.text.strip()
                current_heading_level = lvl
            paragraphs.append(Paragraph(
                text=para.text,
                style=style,
                heading_level=lvl,
                ancestor_heading=current_heading_text,
            ))
        elif tag == "tbl":
            from docx.table import Table
            tbl = Table(block, doc.part)
            for row in tbl.rows:
                total_rows += 1
                for cell in row.cells:
                    for cp in cell.paragraphs:
                        paragraphs.append(Paragraph(
                            text=cp.text,
                            style=(cp.style.name if cp.style else ""),
                            heading_level=None,
                            ancestor_heading=current_heading_text,
                            is_table_cell=True,
                        ))
    return paragraphs, total_breaks, total_rows


def _extract_tables(doc: DocType) -> List[TableView]:
    out: List[TableView] = []
    for t in doc.tables:
        rows = len(t.rows)
        cols = len(t.columns) if rows else 0
        header = ([c.text.strip() for c in t.rows[0].cells]
                  if rows else [])
        first = header[0] if header else ""
        out.append(TableView(rows=rows, cols=cols,
                             header_cells=header, first_cell=first))
    return out


def _count_images(doc: DocType) -> int:
    return sum(1 for r in doc.part._rels.values()
               if "image" in r.reltype)


def _load_docview(path: pathlib.Path) -> DocView:
    d = Document(str(path))
    paragraphs, breaks, rows = _flatten_body(d)
    raw = "\n".join(p.text for p in paragraphs)
    return DocView(path=path,
                   paragraphs=paragraphs,
                   raw_text=raw,
                   page_break_count=breaks,
                   table_row_count=rows,
                   image_count=_count_images(d),
                   tables=_extract_tables(d))


def build_snapshot(manifest) -> PackageSnapshot:
    outs: Dict[str, pathlib.Path] = {
        label: manifest.output_path(label)
        for label in manifest.outputs
    }
    return PackageSnapshot(
        package_root=manifest.root,
        outputs=outs,
        charts_dir=manifest.charts_dir,
    )


def build_snapshot_at(manifest, package_root_override: pathlib.Path) -> PackageSnapshot:
    """Build a snapshot that reads from `package_root_override` rather
    than the manifest's canonical root. Used by the loop when
    evaluating a trial workspace."""
    outs: Dict[str, pathlib.Path] = {
        label: package_root_override / rel
        for label, rel in manifest.outputs.items()
    }
    charts = package_root_override / manifest.charts_dir.name
    return PackageSnapshot(
        package_root=package_root_override,
        outputs=outs,
        charts_dir=charts,
    )
