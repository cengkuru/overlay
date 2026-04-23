"""Deterministic evaluators scoring Overlay review packages.

Every evaluator returns (score, findings). Score is in [0, 1]; findings is a
list of dicts with keys "kind" (the mutation class that might fix it),
"detail" (human-readable), "target_file" (relative path), "target_text"
(a string literal the mutator can search for, if applicable).

The evaluators do not know about mutations. The mutator reads findings and
decides what patch to try.
"""
from __future__ import annotations
import pathlib
import re
import subprocess
import sys
from typing import Callable, Dict, List, Tuple

from docx import Document


# ─── Target rubric: weights and pass thresholds ────────────────────────────

EVALUATOR_WEIGHTS: Dict[str, float] = {
    "E01_no_dashes": 1.0,
    "E02_no_internal_vocab": 1.0,
    "E03_no_ai_tells": 1.0,
    "E04_docs_valid": 1.0,
    "E05_r_ref_integrity": 1.0,
    "E06_no_1480_lead": 1.0,
    "E07_chart_annotations": 0.7,
    "E08_page_length": 0.7,
    "E09_filename_refs": 1.0,
}

PASS_THRESHOLD = 0.95  # per-evaluator target
OVERALL_TARGET = 0.97  # weighted-mean target to declare success


Finding = Dict[str, str]


# ─── Utilities ─────────────────────────────────────────────────────────────

def _extract_text(docx_path: pathlib.Path) -> str:
    """Return concatenated text content from a docx file."""
    d = Document(str(docx_path))
    chunks = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    chunks.append(p.text)
    return "\n".join(chunks)


def _count_images(docx_path: pathlib.Path) -> int:
    d = Document(str(docx_path))
    n = 0
    for rel in d.part.rels.values():
        if "image" in rel.reltype:
            n += 1
    return n


def _page_count_approx(docx_path: pathlib.Path) -> int:
    """Approximate page count from paragraph+table row volume.

    Deterministic but rough. Better than nothing without a Word runtime.
    """
    d = Document(str(docx_path))
    # Count explicit page breaks
    breaks = 0
    for p in d.paragraphs:
        for r in p.runs:
            for br in r._element.iter():
                if br.tag.endswith("}br") and br.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                ) == "page":
                    breaks += 1
    # Every add_page_break in python-docx adds a <w:br w:type="page"/>
    return max(1, breaks + 1)


# ─── Evaluators ────────────────────────────────────────────────────────────

def e01_no_dashes(package: Dict[str, pathlib.Path]) -> Tuple[float, List[Finding]]:
    """No em or en dashes in either document."""
    findings: List[Finding] = []
    total_em = total_en = 0
    for name, path in package.items():
        if not str(path).endswith(".docx"):
            continue
        text = _extract_text(path)
        em = text.count("—")
        en = text.count("–")
        total_em += em
        total_en += en
        if em or en:
            findings.append({
                "kind": "REPLACE_DASH",
                "detail": f"{path.name}: em={em}, en={en}",
                "target_file": "",  # mutator scans build scripts
                "target_text": "",
            })
    score = 1.0 if (total_em == 0 and total_en == 0) else 0.0
    return score, findings


INTERNAL_VOCAB = [
    "Overlay",
    "Gate score",
    "gate item",
    "Reviewer Gate Check",
    "Full Rubric",
    "Part A",
    "Part B",
    "Part C",
]


def e02_no_internal_vocab(package: Dict[str, pathlib.Path]) -> Tuple[float, List[Finding]]:
    findings: List[Finding] = []
    total = 0
    for path in package.values():
        if not str(path).endswith(".docx"):
            continue
        text = _extract_text(path)
        for term in INTERNAL_VOCAB:
            hits = len(re.findall(rf"\b{re.escape(term)}\b", text, re.I))
            if hits:
                total += hits
                findings.append({
                    "kind": "SUB_FORBIDDEN",
                    "detail": f"{path.name}: '{term}' x{hits}",
                    "target_text": term,
                    "target_file": "",
                })
        # Cxx section codes (C0..C12) as standalone words
        cxx = re.findall(r"\bC\d{1,2}\b", text)
        if cxx:
            total += len(cxx)
            findings.append({
                "kind": "SUB_FORBIDDEN",
                "detail": f"{path.name}: section codes {set(cxx)}",
                "target_text": ",".join(set(cxx)),
                "target_file": "",
            })
    score = 1.0 if total == 0 else max(0.0, 1.0 - 0.1 * total)
    return score, findings


AI_TELLS = [
    "delve",
    "leverage ",
    "utiliz",
    "seamless",
    "cutting-edge",
    "groundbreaking",
    "paradigm shift",
    "synergy",
    "holistic",
    "transformative",
    "fostering",
]


def e03_no_ai_tells(package: Dict[str, pathlib.Path]) -> Tuple[float, List[Finding]]:
    findings: List[Finding] = []
    total = 0
    for path in package.values():
        if not str(path).endswith(".docx"):
            continue
        text = _extract_text(path)
        for term in AI_TELLS:
            hits = len(re.findall(rf"\b{re.escape(term)}", text, re.I))
            if hits:
                total += hits
                findings.append({
                    "kind": "DROP_AI_TELL",
                    "detail": f"{path.name}: '{term}' x{hits}",
                    "target_text": term.strip(),
                    "target_file": "",
                })
    score = 1.0 if total == 0 else max(0.0, 1.0 - 0.15 * total)
    return score, findings


def e04_docs_valid(package: Dict[str, pathlib.Path]) -> Tuple[float, List[Finding]]:
    findings: List[Finding] = []
    ok = 0
    expected = 2
    for label in ("review_and_reference", "sample_final_report"):
        path = package.get(label)
        if path is None or not path.exists() or path.stat().st_size < 10_000:
            findings.append({
                "kind": "REBUILD",
                "detail": f"missing or truncated: {label}",
                "target_file": "",
                "target_text": "",
            })
            continue
        try:
            Document(str(path))
            ok += 1
        except Exception as e:
            findings.append({
                "kind": "REBUILD",
                "detail": f"{label} failed to open: {e}",
                "target_file": "",
                "target_text": "",
            })
    return ok / expected, findings


def e05_r_ref_integrity(package: Dict[str, pathlib.Path]) -> Tuple[float, List[Finding]]:
    """Every R-number cited in the merged doc should have a row in its table."""
    findings: List[Finding] = []
    path = package.get("review_and_reference")
    if path is None or not path.exists():
        return 0.0, [{"kind": "REBUILD", "detail": "review_and_reference missing",
                      "target_file": "", "target_text": ""}]
    text = _extract_text(path)
    cited = set(re.findall(r"\bR(?:10|[1-9])\b", text))
    expected = {f"R{i}" for i in range(1, 11)}
    missing = expected - cited
    if missing:
        for m in sorted(missing):
            findings.append({
                "kind": "ADD_RNUMBER",
                "detail": f"R-ref missing from merged doc: {m}",
                "target_text": m,
                "target_file": "",
            })
        return max(0.0, 1.0 - 0.1 * len(missing)), findings
    return 1.0, findings


def e06_no_1480_lead(package: Dict[str, pathlib.Path]) -> Tuple[float, List[Finding]]:
    """The 1,480-slot headline must not appear in executive-summary contexts.

    Soft heuristic: we flag any '1,480' or '1480' reference at all. The
    rubric permits it in an annex but not in the docx outputs.
    """
    findings: List[Finding] = []
    total = 0
    for path in package.values():
        if not str(path).endswith(".docx"):
            continue
        text = _extract_text(path)
        hits = len(re.findall(r"1[,]?480", text))
        if hits:
            total += hits
            findings.append({
                "kind": "REMOVE_1480",
                "detail": f"{path.name}: '1,480' x{hits}",
                "target_text": "1,480",
                "target_file": "",
            })
    score = 1.0 if total == 0 else max(0.0, 1.0 - 0.2 * total)
    return score, findings


def e07_chart_annotations(package: Dict[str, pathlib.Path]) -> Tuple[float, List[Finding]]:
    """Every chart PNG used must exist AND carry annotation signal.

    Deterministic proxy: inspect PNG byte length. Annotated charts are
    significantly larger than bare-bar charts. This is a soft check.
    """
    findings: List[Finding] = []
    charts_dir = package.get("charts_dir")
    if charts_dir is None or not charts_dir.exists():
        return 0.0, [{"kind": "REBUILD", "detail": "charts dir missing",
                      "target_file": "", "target_text": ""}]
    pngs = sorted(charts_dir.glob("0*.png"))
    if not pngs:
        return 0.0, [{"kind": "REBUILD", "detail": "no charts found",
                      "target_file": "", "target_text": ""}]
    small = 0
    for p in pngs:
        if p.stat().st_size < 15_000:  # very small → likely missing annotation
            small += 1
            findings.append({
                "kind": "ANNOTATE_CHART",
                "detail": f"{p.name} is small ({p.stat().st_size} bytes); may lack annotation",
                "target_text": p.name,
                "target_file": "samples/zambia-2026-04/build/build_charts.py",
            })
    score = 1.0 - (small / len(pngs)) if pngs else 0.0
    return score, findings


def e08_page_length(package: Dict[str, pathlib.Path]) -> Tuple[float, List[Finding]]:
    """Page count within target: merged 10-20, final 10-25."""
    findings: List[Finding] = []
    targets = {
        "review_and_reference": (10, 20),
        "sample_final_report": (10, 25),
    }
    scores = []
    for label, (lo, hi) in targets.items():
        path = package.get(label)
        if path is None or not path.exists():
            scores.append(0.0)
            continue
        n = _page_count_approx(path)
        if lo <= n <= hi:
            scores.append(1.0)
        else:
            # distance-based penalty
            dist = min(abs(n - lo), abs(n - hi))
            scores.append(max(0.0, 1.0 - 0.1 * dist))
            findings.append({
                "kind": "ADJUST_LENGTH",
                "detail": f"{label}: estimated pages={n}, target {lo}-{hi}",
                "target_file": "",
                "target_text": label,
            })
    return sum(scores) / len(scores) if scores else 0.0, findings


def e09_filename_refs(package: Dict[str, pathlib.Path]) -> Tuple[float, List[Finding]]:
    """Cross-document filename references must be current.

    The merged doc must reference '02-sample-final-report.docx' (not 03-*).
    """
    findings: List[Finding] = []
    path = package.get("review_and_reference")
    if path is None or not path.exists():
        return 0.0, [{"kind": "REBUILD", "detail": "review_and_reference missing",
                      "target_file": "", "target_text": ""}]
    text = _extract_text(path)
    # Stale reference
    stale = len(re.findall(r"03-sample-final-report\.docx", text))
    if stale:
        findings.append({
            "kind": "FIX_FILENAME_REF",
            "detail": f"stale reference '03-sample-final-report.docx' x{stale}",
            "target_text": "03-sample-final-report.docx",
            "target_file": "samples/zambia-2026-04/build/build_review_and_reference.py",
        })
        return 0.0, findings
    # Must reference the current name
    if "02-sample-final-report.docx" not in text:
        findings.append({
            "kind": "FIX_FILENAME_REF",
            "detail": "current filename reference missing from merged doc",
            "target_text": "02-sample-final-report.docx",
            "target_file": "samples/zambia-2026-04/build/build_review_and_reference.py",
        })
        return 0.5, findings
    return 1.0, findings


# ─── Runner ────────────────────────────────────────────────────────────────

ALL_EVALUATORS: Dict[str, Callable] = {
    "E01_no_dashes": e01_no_dashes,
    "E02_no_internal_vocab": e02_no_internal_vocab,
    "E03_no_ai_tells": e03_no_ai_tells,
    "E04_docs_valid": e04_docs_valid,
    "E05_r_ref_integrity": e05_r_ref_integrity,
    "E06_no_1480_lead": e06_no_1480_lead,
    "E07_chart_annotations": e07_chart_annotations,
    "E08_page_length": e08_page_length,
    "E09_filename_refs": e09_filename_refs,
}


def run_all(package: Dict[str, pathlib.Path]) -> Dict[str, Dict]:
    """Run every evaluator and return a dict of results."""
    results = {}
    for eid, fn in ALL_EVALUATORS.items():
        try:
            score, findings = fn(package)
        except Exception as e:
            score, findings = 0.0, [{"kind": "EVAL_ERROR",
                                     "detail": f"{eid} crashed: {e}",
                                     "target_file": "", "target_text": ""}]
        results[eid] = {
            "score": round(score, 4),
            "weight": EVALUATOR_WEIGHTS[eid],
            "findings": findings,
        }
    return results


def overall_score(results: Dict[str, Dict]) -> float:
    total_weight = sum(r["weight"] for r in results.values())
    weighted = sum(r["score"] * r["weight"] for r in results.values())
    return round(weighted / total_weight, 4)


def package_for_zambia(overlay_root: pathlib.Path) -> Dict[str, pathlib.Path]:
    base = overlay_root / "samples" / "zambia-2026-04"
    return {
        "review_and_reference": base / "01-review-and-reference.docx",
        "sample_final_report": base / "02-sample-final-report.docx",
        "charts_dir": base / "charts",
        "build_dir": base / "build",
        "overlay_root": overlay_root,
    }


if __name__ == "__main__":
    root = pathlib.Path(__file__).resolve().parent.parent
    pkg = package_for_zambia(root)
    res = run_all(pkg)
    print("=" * 60)
    print(f"Overall: {overall_score(res):.4f}")
    print("=" * 60)
    for eid, r in res.items():
        print(f"{eid:30s}  {r['score']:.3f}  (w={r['weight']})")
        for f in r["findings"][:3]:
            print(f"    - {f['detail']}")
