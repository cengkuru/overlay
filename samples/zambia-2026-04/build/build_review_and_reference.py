"""Build Doc 1 (v0.12): CoST IS review + structural reference, merged.

Replaces the separate review letter and structural reference documents.
Pairs REVIEW NOTE callouts (CoST Red) with EXAMPLE callouts (CoST Blue),
section by section, so the country team reads each problem alongside the
pattern the revised report should follow.

Layout: Portrait A4, 22mm margins. Target length 12 to 14 pages.
"""
import pathlib
from docx import Document
from docx.shared import Cm, Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import _docx_helpers as H
from visuals.infographics import cover_photo_strip

CHARTS = pathlib.Path(__file__).parent.parent / "charts"
OUT = pathlib.Path(__file__).parent.parent / "01-review-and-reference.docx"
CONTENT_W = 16.6

# Build the cover photo strip as one PNG at exact content width so it aligns
# flush with the red title block below.
_COVER_STRIP = CHARTS / "cover-strip-review-and-reference.png"
cover_photo_strip(
    [
        H.ASSETS / "stock-africa-rural-road.jpeg",
        H.ASSETS / "stock-africa-highway-aerial.jpeg",
        H.ASSETS / "stock-africa-scaffolding.jpeg",
    ],
    _COVER_STRIP,
    target_width_cm=CONTENT_W,
    strip_height_mm=46,
)


# ─── Document setup ─────────────────────────────────────────────────────────
doc = Document()
H.setup_portrait(doc.sections[0], margin_mm=22)
H.setup_base_style(doc)
H.add_logo_header(doc.sections[0])
H.add_stripe_footer(doc.sections[0])


# ─── Cover ──────────────────────────────────────────────────────────────────
# Single 2-row table: photo strip (row 1) + red title block (row 2). Rows share
# the table's width exactly — pixel-perfect edge alignment by construction.
cover = doc.add_table(rows=2, cols=1)
H.set_table_full_width(cover, CONTENT_W)
H.set_table_zero_cell_margins(cover)
H.apply_col_widths(cover, [CONTENT_W])

# Row 1 — photo strip
scell = cover.rows[0].cells[0]
H.set_cell_margins(scell, top=0, bottom=0, left=0, right=0)
H.set_cell_borders(scell)
scp = scell.paragraphs[0]
scp.paragraph_format.space_before = Pt(0)
scp.paragraph_format.space_after = Pt(0)
scp.paragraph_format.left_indent = Cm(0)
scp.paragraph_format.right_indent = Cm(0)
scp.paragraph_format.first_line_indent = Cm(0)
scp.alignment = WD_ALIGN_PARAGRAPH.LEFT
scp.add_run().add_picture(str(_COVER_STRIP), width=Cm(CONTENT_W))

# Row 2 — red title block
rcell = cover.rows[1].cells[0]
H.set_cell_shading(rcell, H.COVER_RED)
H.set_cell_margins(rcell, top=500, bottom=500, left=480, right=480)
H.set_cell_borders(rcell)
rcp = rcell.paragraphs[0]
rcp.paragraph_format.space_after = Pt(4)
H.add_run(rcp, "CoST IS REVIEW AND STRUCTURAL REFERENCE",
          bold=True, size_pt=11, color=H.WHITE)
rcp2 = rcell.add_paragraph()
rcp2.paragraph_format.space_after = Pt(8)
H.add_run(rcp2, "Feedback on the Zambia OC4IDS field-level mapping report, paired with worked examples of what a revised section should look like.",
          bold=True, size_pt=16, color=H.WHITE)
rcp3 = rcell.add_paragraph()
rcp3.paragraph_format.space_after = Pt(0)
H.add_run(rcp3, "Zambia: ZPPA, NCC, CoST Zambia.  Draft of 3 March 2026 reviewed on 23 April 2026.",
          size_pt=10, color=H.WHITE)

H.para(doc, "", space_after=10)

# How to read (one short block only)
H.callout_box(
    doc, CONTENT_W,
    label="HOW TO READ THIS REVIEW",
    label_color=H.DMUTED,
    body="Short. Section 1 names what is strong in the current draft. Section 2 lists the revisions needed, ordered by priority. Section 3 walks through the revised report section by section. Section 4 sets the next step. This review focuses on refinements, not faults.",
)
H.para(doc, "", space_after=10)

# Metadata
meta = doc.add_table(rows=3, cols=2)
H.set_table_full_width(meta, CONTENT_W)
H.apply_col_widths(meta, [5.0, CONTENT_W - 5.0])
for i, (k, v) in enumerate([
    ("Reviewed by:",
     "CoST International Secretariat. This review applies CoST IS's OC4IDS mapping review methodology."),
    ("Review date:", "23 April 2026."),
    ("Companion:",
     "02-sample-final-report.docx. A fully-worked model report for Zambia. Structure and format are illustrative; stakeholder names are placeholders."),
]):
    kc = meta.rows[i].cells[0]
    vc = meta.rows[i].cells[1]
    for c in (kc, vc):
        H.set_cell_borders(c)
        H.set_cell_margins(c, top=140, bottom=140, left=180, right=180)
    H.add_run(kc.paragraphs[0], k, bold=True, size_pt=9.5, color=H.DMUTED)
    H.add_run(vc.paragraphs[0], v, size_pt=9.5, color=H.CHARCOAL)


# ─── Page 2: Overall assessment + What is strong ───────────────────────────
doc.add_page_break()

H.heading(doc, "Assessment", level=1)
H.para(doc, "The Zambia mapping template is substantively strong: 73 OC4IDS fields mapped, every mapped field backed by a real source path, and 47 of 48 source elements traced forward into OC4IDS. The narrative report does not yet reflect that strength. The executive summary leads with adjectives rather than numbers, lifecycle findings are described but not quantified, and most recommendations have no named owner. The legal framework analysis in Annex 8.1 is publication-quality and its core implication should be pulled into the executive summary; the full legal text stays in the annex. This review asks for ten revisions; most are format and presentation changes, not new analysis.",
       size_pt=10.5, space_after=14)

H.heading(doc, "1.  What is strong", level=1)

H.para(doc, "These elements of the current draft are at the level a reader needs. Keep them on revision.",
       size_pt=10.5, space_after=8)

strengths = [
    ("Legal framework analysis in Annex 8.1.",
     "The citations to the Access to Information Act 2023 (s.6, s.8, s.9, s.17), the ZPPA Act 2020 (s.67, s.70), and the NCC Act 2020 (s.5, s.31, s.33, s.53) are precise and directly tied to disclosure obligations. This is the report's strongest asset. Move the substance forward into the executive summary; keep the full text in the annex."),
    ("The NCC insight.",
     "The observation in section 4.1 that most implementation-stage data is already collected by NCC under its statutory monitoring mandate is the single most actionable finding in the report. It reframes the question from ‘how do we collect this data’ to ‘how do we integrate and disclose what is already collected’. Make this the headline of the recommendations section."),
    ("Source element forward-mapping.",
     "The template maps 47 of the 48 source data elements from the ZPPA e-GP system into OC4IDS paths. That coverage is excellent. Report the number in the methodology section."),
    ("Lifecycle framework in section 8.2.",
     "The five-row data-source-by-stage table is the foundation for the decision summary this review asks for. Expand it into a four-bucket classification (ready to publish, small format fixes, needs system work, needs policy action)."),
    ("Mapping quality.",
     "Every OC4IDS path the template lists as mapped contains real source content. No placeholders. A reviewer sampling the mapping sheets can verify every line back to the ZPPA e-GP schema."),
]
for title, body in strengths:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    H.add_run(p, "+  ", bold=True, size_pt=11, color=H.LOW_TEXT)
    H.add_run(p, title + "  ", bold=True, size_pt=10.5, color=H.CHARCOAL)
    H.add_run(p, body, size_pt=10.5, color=H.CHARCOAL)


# ─── Section 2: Priority revisions (promoted to the top) ───────────────────
doc.add_page_break()
H.heading(doc, "2.  Priority revisions", level=1)

H.para(doc, "Ten revisions, ordered by priority. Critical items close the biggest readability gaps. High items sharpen the structure. Medium and low items are polish. Proposed owners are placeholders; the MSG sets final assignments and sequencing.",
       size_pt=10.5, space_after=10)

revs = [
    ("R1", "Critical",
     "Executive summary carries no number.",
     "Rewrite the opening paragraph to lead with the headline numbers (73 fields mapped, 2 of 4 sheets near zero) and the two-sentence structural reading.",
     "Zambia Technical Team"),
    ("R2", "Critical",
     "Phase-level percentages not reported.",
     "Add a phase coverage table to section 4: one row per phase with the fraction, the percentage, the dominant cause, and one sentence of evidence.",
     "Zambia Technical Team"),
    ("R4", "Critical",
     "Recommendations section is out of scope for a mapping report.",
     "Replace the current section 6.1 to 6.6 recommendations block with three directional moves only: what can be published now, what needs a short format-fix pass, and what needs the NCC and ZPPA integration workstream. Detailed recommendations with named owners, target quarters, budgets, and a legal crosswalk are a separate reform-planning product, to be developed after MSG approval of this mapping.",
     "Zambia Technical Team"),
    ("R3", "High",
     "Where each field comes from is not shown.",
     "Add a source column to the section 4 or 5 tables, or include a short source-paths annex. One row per mapped field.",
     "Zambia Technical Team"),
    ("R5", "High",
     "Gaps not classified by cause.",
     "Apply the five-cause typology (not collected, collected not disclosed, unstructured, inconsistent, restricted by law) to every material gap. Label the NCC insight as ‘collected but not yet disclosed’.",
     "Zambia Technical Team"),
    ("R6", "High",
     "Data beyond OC4IDS not examined.",
     "Add a short section listing non-OC4IDS fields the country publishes. Recommend which to preserve as local extensions and which to propose upstream to OCP.",
     "Zambia Technical Team with CoST IS"),
    ("R7", "High",
     "Conclusion does not classify what can be published now.",
     "Add a four-bucket decision summary at the end: ready to publish, small format fixes, needs system work, needs policy action. This is the single view a senior reader scans first.",
     "Zambia Technical Team"),
    ("R8", "Medium",
     "Portal URL and access dates missing from the body.",
     "Add the ZPPA e-GP URL and the dates accessed to section 3.1.",
     "Zambia Technical Team"),
    ("R9", "Medium",
     "Sample not quantified.",
     "State count, aggregate value, sectors, time window, and selection method in section 3.",
     "Zambia Technical Team"),
    ("R10", "Low",
     "Unmapped template rows lack reasons.",
     "Before resubmission, add a short Notes entry to every unmapped row in the OC4IDS sheets.",
     "Zambia Technical Team"),
]

prio_color = {
    "Critical": H.RED,
    "High":     H.BLUE,
    "Medium":   H.MEDIUM_TEXT,
    "Low":      H.LOW_TEXT,
}

# Group items by priority so the tier label appears once per group, not per item.
from itertools import groupby
tier_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
revs_sorted = sorted(revs, key=lambda r: tier_order[r[1]])

for tier, items_iter in groupby(revs_sorted, key=lambda r: r[1]):
    # Tier heading (one row per tier, not per item)
    th = doc.add_paragraph()
    th.paragraph_format.space_before = Pt(6)
    th.paragraph_format.space_after = Pt(4)
    H.add_run(th, tier.upper(), bold=True, size_pt=11, color=prio_color[tier])

    for ref, _, issue, rec, owner in items_iter:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        H.add_run(p, f"{ref}  ", bold=True, size_pt=10.5, color=H.CHARCOAL)
        H.add_run(p, issue + "  ", bold=True, size_pt=10.5, color=H.CHARCOAL)
        H.add_run(p, rec, size_pt=10.5, color=H.CHARCOAL)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(10)
        p2.paragraph_format.left_indent = Cm(0.6)
        H.add_run(p2, f"Proposed owner: {owner}.", size_pt=9.5, color=H.DMUTED, italic=True)


# ─── Section 3: Section-by-section guidance ────────────────────────────────
doc.add_page_break()
H.heading(doc, "3.  Section-by-section guidance", level=1)

H.para(doc, "For each major section of the revised report, this is what the current draft does and what the revision should do instead. Quoted lines are from the draft of 3 March 2026, punctuation normalised; the Zambia team may refine them against the final text. R4 is cross-cutting (it asks for the current recommendations block to be replaced with three directional moves) and is covered in Section 2 only.",
       size_pt=10.5, space_after=10)

# 3.1 Executive summary
H.heading(doc, "3.1  Executive summary (addresses R1)", level=2)
H.quoted_passage(
    doc, CONTENT_W,
    source_hint="Executive summary, para 1 (draft of 3 March 2026)",
    body="The report says: “Significant data gaps exist in Zambia’s disclosure of infrastructure procurement and implementation data across the OC4IDS standard.”",
)
H.para(doc, "", space_after=6)
H.para(doc, "Revision: lead with three numbers, not three adjectives. Open with how many OC4IDS fields are currently mapped (73), how many more become publishable within twelve months, and whether existing law already authorises publication of most of what is missing (it does). Adjectives follow; numbers lead.",
       size_pt=10.5, space_after=14)

# 3.2 Findings by lifecycle phase
H.heading(doc, "3.2  Findings by lifecycle phase (addresses R2)", level=2)
H.quoted_passage(
    doc, CONTENT_W,
    source_hint="Section 4, opening paragraph (draft of 3 March 2026)",
    body="The report says: “Implementation-stage disclosure is thin, with most fields unpopulated or populated only in narrative form. Maintenance and decommissioning remain largely unaddressed by the current template submission.”",
)
H.para(doc, "", space_after=6)
H.para(doc, "Revision: replace the qualitative narrative with a quantified phase table. One row per phase with fraction, percentage, dominant cause, and one sentence of evidence. Implementation is the most important row; label it ‘collected but not yet publicly disclosed’ so readers see the integration-not-collection framing immediately.",
       size_pt=10.5, space_after=14)

# 3.3 Where each field comes from
H.heading(doc, "3.3  Where each field comes from (addresses R3)", level=2)
H.para(doc, "The template already traces every mapped OC4IDS path to a specific ZPPA e-GP source path. The report does not surface this. A reader of the narrative cannot see what system each field comes from.",
       size_pt=10.5, space_after=6)
H.para(doc, "Revision: add a source column to the phase tables, or include a short source-paths annex. Each row carries the OC4IDS path, the originating ZPPA e-GP field name, the population status, and the disclosure pathway.",
       size_pt=10.5, space_after=14)

# 3.4 Why the gaps exist
H.heading(doc, "3.4  Why the gaps exist (addresses R5)", level=2)
H.quoted_passage(
    doc, CONTENT_W,
    source_hint="Section 4.1 or equivalent, on implementation data (draft of 3 March 2026)",
    body="The report says: “The National Council for Construction maintains inspection and project monitoring records under its statutory mandate, although these are not currently integrated with ZPPA’s disclosure channels.”",
)
H.para(doc, "", space_after=6)
H.para(doc, "Revision: classify every material gap against one of five causes (not collected, collected but not yet disclosed, only in unstructured documents, collected inconsistently, or restricted by law or workflow). Each cause takes a different fix. The Zambia-dominant cause is ‘collected but not yet disclosed’. Labelling it that way reframes the whole recommendations section.",
       size_pt=10.5, space_after=14)

# 3.5 Data beyond OC4IDS
H.heading(doc, "3.5  Data beyond OC4IDS (addresses R6)", level=2)
H.para(doc, "The current draft does not examine what the ZPPA e-GP or NCC systems publish that falls outside OC4IDS. Country-specific disclosure that could feed upstream to OCP is invisible. Preserving useful local fields as extensions is part of a mature mapping report.",
       size_pt=10.5, space_after=6)
H.para(doc, "Revision: add a short section listing country-specific fields worth preserving as local extensions, with one or two worth proposing upstream. Each entry names the field and the reason.",
       size_pt=10.5, space_after=14)

# 3.6 Decision summary
H.heading(doc, "3.6  Decision summary (addresses R7)", level=2)
H.quoted_passage(
    doc, CONTENT_W,
    source_hint="Section 8.2 (draft of 3 March 2026), data-holder lifecycle table",
    body="The report says (section 8.2, row on implementation): “Implementation, NCC, Project monitoring records, inspection forms, progress reports.”",
)
H.para(doc, "", space_after=6)
H.para(doc, "Revision: close the report with a four-bucket decision summary. Classify every publishable element into ready to publish, small format fixes, needs system work, or needs policy action. Each pathway carries a field count and one or two examples. This is the single view a senior reader scans first.",
       size_pt=10.5, space_after=14)

# Closing reference to the sample report — one clear pointer, not scattered.
H.callout_box(
    doc, CONTENT_W,
    label="WORKED EXAMPLES",
    label_color=H.DMUTED,
    body="The companion file, 02-sample-final-report.docx, carries the same sections filled out as a worked model for Zambia, including the phase table, source-paths annex, five-cause typology, and four-bucket decision summary described above. Use it as a reference to calibrate the revised report, not as a template to copy; stakeholder names in the sample are placeholders pending MSG consultation.",
)


# ─── Section 4: Acceptance threshold and next step ─────────────────────────
doc.add_page_break()
H.heading(doc, "4.  Acceptance threshold and next step", level=1)

H.heading(doc, "4.1  When the revision is ready", level=2)
H.para(doc, "The revised report is acceptable for CoST IS countersign when all three Critical items (R1, R2, R4) are closed and each High item (R3, R5, R6, R7) is either closed or documented as out of scope with brief reasoning. Medium and low items can follow in a subsequent polish pass.",
       size_pt=10.5, space_after=14)

H.heading(doc, "4.2  Next step", level=2)
H.para(doc, "A short call between CoST IS and the Zambia Technical Team to walk through the ten revisions and agree a resubmission schedule with the MSG. Critical items should land first, High items in a second pass, Medium and Low in polish. The MSG sets the schedule.",
       size_pt=10.5, space_after=14)

H.heading(doc, "4.3  What the revision unlocks", level=2)
H.para(doc, "Once these revisions land, the report is fit for CoST IS countersign and member-country use. Anything beyond that (scoping an integration platform, designating ATI officers, drafting a data-sharing instrument) belongs to a separate planning conversation.",
       size_pt=10.5, space_after=14)

H.para(doc, "Source: Zambia OC4IDS field-level mapping template v0.9.5, accessed 2026-04-23. Zambia OC4IDS field-level mapping report (draft), 3 March 2026.",
       size_pt=9, italic=True, color=H.MUTED, space_after=0)

doc.save(str(OUT))
print(f"OK {OUT}")
