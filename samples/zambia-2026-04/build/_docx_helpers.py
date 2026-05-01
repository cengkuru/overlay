"""Shared docx helpers for all three Zambia deliverables.

Consolidates the table-width, cell-shading, border, padding, and image-centering
code that was duplicated across three build scripts.

Design goals for v0.11:
- Generous inner padding so table content does not hug cell edges.
- Tables always span the section content width unless explicitly narrower.
- Images centered by default with breathing room above and below.
- Single source of truth for the CoST palette.
"""
from __future__ import annotations
import pathlib
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── CoST palette ───────────────────────────────────────────────────────────
RED = RGBColor(0xB7, 0x25, 0x1C)
COVER_RED = RGBColor(0xCC, 0x20, 0x28)
TABLE_RED = RGBColor(0xC0, 0x00, 0x00)
CHARCOAL = RGBColor(0x4A, 0x47, 0x43)
BLUE = RGBColor(0x71, 0xB6, 0xC9)
YELLOW = RGBColor(0xF1, 0xC0, 0x3D)
MUTED = RGBColor(0x9C, 0x9B, 0x99)
DMUTED = RGBColor(0x6B, 0x6A, 0x68)
LMUTED = RGBColor(0xB0, 0xB0, 0xB1)
MGRAY = RGBColor(0xE5, 0xE4, 0xE3)
LGRAY = RGBColor(0xF3, 0xF2, 0xF1)
LIGHT_NEUTRAL = RGBColor(0xF0, 0xED, 0xE7)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LOW_TEXT = RGBColor(0x4E, 0xA8, 0x3D)

# Priority tints (review letter)
CRITICAL_BG = RGBColor(0xF5, 0xD5, 0xD3)
HIGH_BG = RGBColor(0xD9, 0xED, 0xF3)
MEDIUM_BG = RGBColor(0xFD, 0xF3, 0xD7)
LOW_BG = RGBColor(0xE8, 0xF5, 0xE9)
MEDIUM_TEXT = RGBColor(0xD4, 0xA0, 0x17)

ASSETS = pathlib.Path("/Users/cengkurumichael/.claude/skills/cost-document-design/assets")

# Generous padding defaults (v0.11 increase from v0.10)
# dxa units: 20 dxa = 1 pt; 567 dxa = 1 cm
PAD_TOP = 180
PAD_BOTTOM = 180
PAD_LEFT = 240
PAD_RIGHT = 240

HEADER_PAD_TOP = 220
HEADER_PAD_BOTTOM = 220

CALLOUT_PAD_TOP = 260
CALLOUT_PAD_BOTTOM = 260
CALLOUT_PAD_LEFT = 300
CALLOUT_PAD_RIGHT = 300


def hex_of(c):
    if isinstance(c, str):
        return c.upper().lstrip("#")
    return "".join(f"{b:02X}" for b in c)


# ─── Cell helpers ───────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_of(color))
    tc_pr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None,
                     color="E5E4E3", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), val if val else 'nil')
        if val:
            border.set(qn('w:sz'), sz)
            border.set(qn('w:color'), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_cell_margins(cell, *, top=PAD_TOP, bottom=PAD_BOTTOM,
                     left=PAD_LEFT, right=PAD_RIGHT):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for k, v in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{k}')
        el.set(qn('w:w'), str(v))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tc_pr.append(tcMar)


def set_vertical_align(cell, align="top"):
    """align: 'top', 'center', 'bottom'."""
    tc_pr = cell._tc.get_or_add_tcPr()
    v = OxmlElement('w:vAlign')
    v.set(qn('w:val'), align)
    tc_pr.append(v)


# ─── Table helpers ──────────────────────────────────────────────────────────

def set_table_full_width(table, content_cm):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn('w:tblW')):
        tbl_pr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(content_cm * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tbl_pr.append(tblW)
    # Force table indent = 0 so the table starts exactly at the section margin.
    for existing in tbl_pr.findall(qn('w:tblInd')):
        tbl_pr.remove(existing)
    tbl_ind = OxmlElement('w:tblInd')
    tbl_ind.set(qn('w:w'), '0')
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_ind)


def set_table_zero_cell_margins(table):
    """Strip table-level default cell margins (tblCellMar). Needed for
    edge-to-edge content like the cover photo strip, where any inherited
    0.19cm inner padding would make a CONTENT_W-wide image overflow.
    """
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn('w:tblCellMar')):
        tbl_pr.remove(existing)
    cell_mar = OxmlElement('w:tblCellMar')
    for side in ('top', 'start', 'bottom', 'end', 'left', 'right'):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        cell_mar.append(m)
    tbl_pr.append(cell_mar)


def apply_col_widths(table, widths_cm):
    tbl = table._tbl
    for grid in tbl.findall(qn('w:tblGrid')):
        tbl.remove(grid)
    grid = OxmlElement('w:tblGrid')
    for w in widths_cm:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w * 567)))
        grid.append(gc)
    tbl.insert(list(tbl).index(tbl.find(qn('w:tblPr'))) + 1, grid)
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def apply_default_padding(table, *, header_row=0):
    """Apply standard padding to every cell. Header row gets taller padding."""
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == header_row:
                set_cell_margins(cell, top=HEADER_PAD_TOP, bottom=HEADER_PAD_BOTTOM,
                                 left=PAD_LEFT, right=PAD_RIGHT)
            else:
                set_cell_margins(cell)
            set_vertical_align(cell, "top")


def set_row_height(row, height_cm, rule="atLeast"):
    """Keep rows from collapsing. rule: 'atLeast' or 'exact'."""
    tr_pr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'), str(int(height_cm * 567)))
    h.set(qn('w:hRule'), rule)
    tr_pr.append(h)


# ─── Paragraph / run helpers ───────────────────────────────────────────────

def add_run(p, text, *, bold=False, italic=False, size_pt=10.5,
            color=CHARCOAL, font="Arial"):
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def para(doc, text, *, size_pt=10.5, color=CHARCOAL, bold=False, italic=False,
         space_after=6, space_before=0, align=None, font="Arial"):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if text:
        add_run(p, text, bold=bold, italic=italic, size_pt=size_pt,
                color=color, font=font)
    return p


def heading(doc, text, *, level=1, color=None):
    sizes = {1: 15, 2: 12, 3: 11}
    space_before = {1: 14, 2: 10, 3: 6}
    c = color or (RED if level == 1 else CHARCOAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before[level])
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, bold=True, size_pt=sizes[level], color=c)
    return p


def accent_bar(doc, color=RED, width_cm=8):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = t._tbl.tblPr
    for existing in tbl_pr.findall(qn('w:tblW')):
        tbl_pr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_cm * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tbl_pr.append(tblW)
    cell = t.cell(0, 0)
    cell.text = ""
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    set_cell_borders(cell, bottom="single", color=hex_of(color), sz="24")
    return t


def centered_image(doc, image_path, width_cm, *, space_before=8, space_after=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))


# ─── Page / section helpers ────────────────────────────────────────────────

def setup_portrait(section, margin_mm=22):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(margin_mm)
    section.bottom_margin = Mm(margin_mm)
    section.left_margin = Mm(margin_mm)
    section.right_margin = Mm(margin_mm)
    return 210 - 2 * margin_mm  # content width in mm


def setup_landscape(section, margin_mm=20):
    from docx.enum.section import WD_ORIENT
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(margin_mm)
    section.bottom_margin = Mm(margin_mm)
    section.left_margin = Mm(margin_mm)
    section.right_margin = Mm(margin_mm)
    return 297 - 2 * margin_mm


def add_logo_header(section, *, logo_width_cm=4.2):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    hp.add_run().add_picture(str(ASSETS / "cost-logo-real.jpeg"), width=Cm(logo_width_cm))


def add_stripe_footer(section, *, stripe_width_cm=17):
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.space_after = Pt(0)
    fp.add_run().add_picture(str(ASSETS / "cost-stripe-real.png"),
                             width=Cm(stripe_width_cm))
    fp2 = footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp2.paragraph_format.space_before = Pt(4)
    fr = fp2.add_run()
    fr.font.name = 'Arial'
    fr.font.size = Pt(8.5)
    fr.font.color.rgb = MUTED
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    fr._element.append(fld)


def setup_base_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    style.font.color.rgb = CHARCOAL


# ─── Common block builders ─────────────────────────────────────────────────

def verdict_box(doc, content_w_cm, *, label, headline, tail=None):
    """Verdict callout: top+bottom red rules, grey fill, padded."""
    vb = doc.add_table(rows=1, cols=1)
    set_table_full_width(vb, content_w_cm)
    apply_col_widths(vb, [content_w_cm])
    vc = vb.cell(0, 0)
    set_cell_shading(vc, LGRAY)
    set_cell_borders(vc, top="single", bottom="single", color=hex_of(RED), sz="24")
    set_cell_margins(vc, top=CALLOUT_PAD_TOP, bottom=CALLOUT_PAD_BOTTOM,
                     left=CALLOUT_PAD_LEFT, right=CALLOUT_PAD_RIGHT)
    p = vc.paragraphs[0]
    add_run(p, label + "   ", bold=True, size_pt=10, color=RED)
    add_run(p, headline, size_pt=10.5, color=CHARCOAL)
    if tail:
        p2 = vc.add_paragraph()
        p2.paragraph_format.space_before = Pt(6)
        add_run(p2, tail, size_pt=10.5, color=CHARCOAL)
    return vb


def callout_box(doc, content_w_cm, *, label, label_color, body):
    """Left-border callout box for advisory notes."""
    cb = doc.add_table(rows=1, cols=1)
    set_table_full_width(cb, content_w_cm)
    apply_col_widths(cb, [content_w_cm])
    cc = cb.cell(0, 0)
    set_cell_shading(cc, LGRAY)
    set_cell_borders(cc, left="single", color=hex_of(label_color), sz="32")
    set_cell_margins(cc, top=CALLOUT_PAD_TOP, bottom=CALLOUT_PAD_BOTTOM,
                     left=CALLOUT_PAD_LEFT, right=CALLOUT_PAD_RIGHT)
    p = cc.paragraphs[0]
    add_run(p, label + "   ", bold=True, size_pt=10, color=label_color)
    add_run(p, body, size_pt=10.5, color=CHARCOAL)
    return cb


def quoted_passage(doc, content_w_cm, *, source_hint, body):
    """QUOTED FROM DRAFT block: muted grey left-border, italic body.

    Renders a verbatim (or near-verbatim) passage from the document under
    review, so the downstream REVIEW NOTE has a specific target. The source
    hint names the section the quote is from (e.g. "Executive summary, para 1"
    or "Section 4.1, opening sentence").

    Visual convention: muted grey (not red, not blue) so the reader sees the
    three-step pattern at a glance — quote (muted) → note (red) → example (blue).
    """
    cb = doc.add_table(rows=1, cols=1)
    set_table_full_width(cb, content_w_cm)
    apply_col_widths(cb, [content_w_cm])
    cc = cb.cell(0, 0)
    set_cell_shading(cc, LGRAY)
    set_cell_borders(cc, left="single", color=hex_of(MUTED), sz="36")
    set_cell_margins(cc, top=CALLOUT_PAD_TOP, bottom=CALLOUT_PAD_BOTTOM,
                     left=CALLOUT_PAD_LEFT, right=CALLOUT_PAD_RIGHT)
    label_p = cc.paragraphs[0]
    label_p.paragraph_format.space_after = Pt(4)
    add_run(label_p, "QUOTED FROM DRAFT", bold=True, size_pt=9.5, color=DMUTED)
    if source_hint:
        add_run(label_p, "   " + source_hint, size_pt=9.5, color=DMUTED, italic=True)
    body_p = cc.add_paragraph()
    body_p.paragraph_format.space_before = Pt(0)
    body_p.paragraph_format.space_after = Pt(0)
    add_run(body_p, body, size_pt=10.5, color=CHARCOAL, italic=True)
    return cb


def review_note(doc, content_w_cm, *, ref_ids, body):
    """REVIEW NOTE callout: CoST Red left-border, grey fill.

    Used to annotate what the current draft does (or fails to do) in each
    section of the merged review-and-reference document.

    ref_ids: list of revision identifiers (e.g. ["R1"]) shown after the label.
    """
    cb = doc.add_table(rows=1, cols=1)
    set_table_full_width(cb, content_w_cm)
    apply_col_widths(cb, [content_w_cm])
    cc = cb.cell(0, 0)
    set_cell_shading(cc, LGRAY)
    set_cell_borders(cc, left="single", color=hex_of(RED), sz="36")
    set_cell_margins(cc, top=CALLOUT_PAD_TOP, bottom=CALLOUT_PAD_BOTTOM,
                     left=CALLOUT_PAD_LEFT, right=CALLOUT_PAD_RIGHT)
    p = cc.paragraphs[0]
    add_run(p, "REVIEW NOTE", bold=True, size_pt=9.5, color=RED)
    refs = "  (" + ", ".join(ref_ids) + ")" if ref_ids else ""
    add_run(p, refs + "   ", bold=True, size_pt=9.5, color=DMUTED)
    add_run(p, body, size_pt=10.5, color=CHARCOAL)
    return cb


def example_block(doc, content_w_cm, *, label_detail=None, content_fn):
    """EXAMPLE block: CoST Blue left-border, white fill. Content rendered via callback.

    The callback receives the inner cell object; use it to insert paragraphs,
    tables, images, or any combination the example requires. A label row
    ("EXAMPLE" + optional detail) is pre-written into the cell.
    """
    eb = doc.add_table(rows=1, cols=1)
    set_table_full_width(eb, content_w_cm)
    apply_col_widths(eb, [content_w_cm])
    ec = eb.cell(0, 0)
    set_cell_shading(ec, WHITE)
    set_cell_borders(ec, left="single", color=hex_of(BLUE), sz="36")
    set_cell_margins(ec, top=CALLOUT_PAD_TOP, bottom=CALLOUT_PAD_BOTTOM,
                     left=CALLOUT_PAD_LEFT, right=CALLOUT_PAD_RIGHT)
    label_p = ec.paragraphs[0]
    label_p.paragraph_format.space_after = Pt(6)
    add_run(label_p, "EXAMPLE", bold=True, size_pt=9.5, color=BLUE)
    if label_detail:
        add_run(label_p, "   " + label_detail, size_pt=9.5, color=DMUTED, italic=True)
    content_fn(ec)
    return eb


def styled_table_header(row, labels, *, bg=TABLE_RED, font_color=WHITE, size_pt=10):
    for i, h in enumerate(labels):
        c = row.cells[i]
        set_cell_shading(c, bg)
        set_cell_borders(c, top="single", bottom="single", left="single",
                         right="single", color=hex_of(bg), sz="4")
        set_cell_margins(c, top=HEADER_PAD_TOP, bottom=HEADER_PAD_BOTTOM,
                         left=PAD_LEFT, right=PAD_RIGHT)
        add_run(c.paragraphs[0], h, bold=True, size_pt=size_pt, color=font_color)
