"""Build Doc 2: a sample-corrected version of Zambia's OC4IDS report.

Archetype: Activity Report. Intensity: Full Publication. Layout: Portrait A4.

Preserves Zambia's real content (legal framework, NCC insight, lifecycle framing)
and rewrites the parts Overlay flagged: quantified headline, phase percentages,
source-field provenance, owned recommendations, decision summary panel.

This is a TARGET-STATE reference. Use it to calibrate the resubmission, not to copy.
"""
import pathlib
from docx import Document
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── CoST palette ───────────────────────────────────────────────────────────
RED = RGBColor(0xB7, 0x25, 0x1C)
COVER_RED = RGBColor(0xCC, 0x20, 0x28)
TABLE_RED = RGBColor(0xC0, 0x00, 0x00)
CHARCOAL = RGBColor(0x4A, 0x47, 0x43)
BLUE = RGBColor(0x71, 0xB6, 0xC9)
YELLOW = RGBColor(0xF1, 0xC0, 0x3D)
MUTED = RGBColor(0x9C, 0x9B, 0x99)
DMUTED = RGBColor(0x6B, 0x6A, 0x68)
LMUTED = RGBColor(0xB0, 0xB0, 0xB1)
MGRAY = RGBColor(0xE5, 0xE4, 0xE3)
LGRAY = RGBColor(0xF3, 0xF2, 0xF1)
LIGHT_NEUTRAL = RGBColor(0xF0, 0xED, 0xE7)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LOW_TEXT = RGBColor(0x4E, 0xA8, 0x3D)

ASSETS = pathlib.Path("/Users/cengkurumichael/.claude/skills/cost-document-design/assets")
CHARTS = pathlib.Path(__file__).parent.parent / "charts"
OUT = pathlib.Path(__file__).parent.parent / "02-sample-corrected-report.docx"


def hex_of(c):
    if isinstance(c, str):
        return c.upper().lstrip("#")
    return "".join(f"{b:02X}" for b in c)


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_of(color))
    tc_pr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, color="E5E4E3", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), val if val else 'nil')
        if val:
            border.set(qn('w:sz'), sz)
            border.set(qn('w:color'), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for k, v in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{k}')
        el.set(qn('w:w'), str(v))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tc_pr.append(tcMar)


def add_run(p, text, *, bold=False, italic=False, size_pt=10.5, color=CHARCOAL, font="Arial"):
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def para(doc, text, *, size_pt=10.5, color=CHARCOAL, bold=False, italic=False,
         space_after=6, align=None, font="Arial"):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if text:
        add_run(p, text, bold=bold, italic=italic, size_pt=size_pt, color=color, font=font)
    return p


def heading(doc, text, *, level=1):
    sizes = {1: 16, 2: 13, 3: 11}
    space_before = {1: 18, 2: 12, 3: 8}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before[level])
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, bold=True, size_pt=sizes[level], color=CHARCOAL if level > 1 else RED)
    if level == 1:
        # thin accent rule
        t = doc.add_table(rows=1, cols=1)
        c = t.cell(0, 0)
        set_cell_borders(c, bottom="single", color=hex_of(RED), sz="24")
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(120 * 20))
        tblW.set(qn('w:type'), 'dxa')
        t._tbl.tblPr.append(tblW)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(6)  # page break


# ─── Build document ─────────────────────────────────────────────────────────

doc = Document()
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = Mm(20)
sec.bottom_margin = Mm(20)
sec.left_margin = Mm(20)
sec.right_margin = Mm(20)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)
style.font.color.rgb = CHARCOAL

# Header: logo
header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hp.paragraph_format.space_after = Pt(0)
hp.add_run().add_picture(str(ASSETS / "cost-logo-real.jpeg"), width=Cm(4.2))

# Footer: stripe + page number
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
fp.paragraph_format.space_after = Pt(0)
fp.add_run().add_picture(str(ASSETS / "cost-stripe-real.png"), width=Cm(17))
fp2 = footer.add_paragraph()
fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fp2.paragraph_format.space_before = Pt(4)
fr = fp2.add_run()
fr.font.name = 'Arial'; fr.font.size = Pt(8.5); fr.font.color.rgb = MUTED
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'PAGE')
fr._element.append(fld)

# ─── Cover block ────────────────────────────────────────────────────────────
# Small stock image strip
strip = doc.add_table(rows=1, cols=3)
strip.autofit = False
widths = [Cm(5.6), Cm(5.6), Cm(5.6)]
stocks = [
    ASSETS / "stock-africa-rural-road.jpeg",
    ASSETS / "stock-site-inspection.jpeg",
    ASSETS / "stock-africa-excavation.jpeg",
]
for i, w in enumerate(widths):
    strip.rows[0].cells[i].width = w
    set_cell_borders(strip.rows[0].cells[i])
    strip.rows[0].cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
    strip.rows[0].cells[i].paragraphs[0].add_run().add_picture(str(stocks[i]), width=w, height=Mm(38))

# Red title block
red = doc.add_table(rows=1, cols=1)
red.autofit = False
red.rows[0].cells[0].width = Cm(17)
rcell = red.rows[0].cells[0]
set_cell_shading(rcell, COVER_RED)
set_cell_margins(rcell, top=380, bottom=380, left=380, right=380)
set_cell_borders(rcell)
rcp = rcell.paragraphs[0]
rcp.paragraph_format.space_after = Pt(4)
add_run(rcp, "OC4IDS FIELD-LEVEL MAPPING REPORT", bold=True, size_pt=12, color=WHITE)
rcp2 = rcell.add_paragraph()
rcp2.paragraph_format.space_after = Pt(4)
add_run(rcp2, "Findings, gaps, and the path to an Information Platform for Public Infrastructure in Zambia",
        bold=True, size_pt=20, color=WHITE)
rcp3 = rcell.add_paragraph()
rcp3.paragraph_format.space_after = Pt(0)
add_run(rcp3, "Sample corrected reference  ·  Overlay v0.9  ·  2026-04-23", size_pt=10, color=WHITE)

para(doc, "", space_after=4)

# Metadata under red block
meta = doc.add_table(rows=4, cols=2)
meta.autofit = False
for r in meta.rows:
    r.cells[0].width = Cm(4)
    r.cells[1].width = Cm(13)
for i, (k, v) in enumerate([
    ("Publisher:", "National Council for Construction (NCC) and Zambia Public Procurement Authority (ZPPA); coordinated by CoST Zambia."),
    ("Prepared by:", "Zambia Technical Team (NCC, ZPPA, CoST Zambia), with support from CoST International Secretariat."),
    ("Review framework:", "OC4IDS v0.9.5 Field-Level Mapping Template. Overlay v0.9 rubric applied 2026-04-23."),
    ("Publication licence:", "CC BY 4.0  ·  www.infrastructuretransparency.org"),
]):
    kc = meta.rows[i].cells[0]
    vc = meta.rows[i].cells[1]
    for c in (kc, vc):
        set_cell_borders(c)
    add_run(kc.paragraphs[0], k, bold=True, size_pt=9.5, color=DMUTED)
    add_run(vc.paragraphs[0], v, size_pt=9.5, color=CHARCOAL)

# Legal footer
para(doc, "CoST, the Infrastructure Transparency Initiative  ·  Company number 8159144  ·  Charity number 1152236",
     size_pt=7.5, italic=True, color=LMUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

# ─── Page 2: Executive summary ──────────────────────────────────────────────
doc.add_page_break()
heading(doc, "Executive summary", level=1)

# Headline verdict box
vb = doc.add_table(rows=1, cols=1)
vc = vb.cell(0, 0)
set_cell_shading(vc, LGRAY)
set_cell_borders(vc, top="single", bottom="single", color=hex_of(RED), sz="24")
set_cell_margins(vc, top=200, bottom=200, left=240, right=240)
vp = vc.paragraphs[0]
add_run(vp, "VERDICT   ", bold=True, size_pt=10, color=RED)
add_run(vp, "Zambia's e-GP system can publish 35 OC4IDS fields today. Another 20 need light transformation. Most of the rest is already collected by the National Council for Construction under the NCC Act 2020; it is a disclosure gap, not a data gap.",
        bold=True, size_pt=11, color=CHARCOAL)
vp2 = vc.add_paragraph()
vp2.paragraph_format.space_before = Pt(4)
add_run(vp2, "The Access to Information Act 2023 (s.8) already mandates proactive publication of 46 of the missing fields. The proposed Information Platform for Public Infrastructure (IPPI-Zambia) is the integration layer this analysis calls for.",
        size_pt=10.5, color=CHARCOAL)

para(doc, "", space_after=10)

heading(doc, "Headline numbers", level=2)

# Stat strip
stats = doc.add_table(rows=1, cols=4)
stats.autofit = False
for c in stats.rows[0].cells:
    c.width = Cm(4.2)

stat_data = [
    ("4.9%", "of the 1,480 OC4IDS template slots mapped from the ZPPA e-GP system (73 fields).", RED),
    ("97.9%", "of the 48 ZPPA e-GP source data elements successfully forward-mapped to OC4IDS.", BLUE),
    ("35", "OC4IDS fields can be published immediately with no new system work.", LOW_TEXT),
    ("46", "OC4IDS fields that the Access to Information Act 2023 already mandates for disclosure.", YELLOW),
]

for i, (val, label, col) in enumerate(stat_data):
    cell = stats.rows[0].cells[i]
    set_cell_shading(cell, LGRAY)
    set_cell_margins(cell, top=200, bottom=200, left=200, right=200)
    set_cell_borders(cell, top="single", color=hex_of(col), sz="32")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p1, val, bold=True, size_pt=22, color=col)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    add_run(p2, label, size_pt=9, color=DMUTED)

para(doc, "", space_after=12)

heading(doc, "Top three strengths", level=2)
for s in [
    "Legal framework analysis already cites ATI Act 2023 s.6/8/9/17, ZPPA Act No.3/2020 s.67/70, and NCC Act No.10/2020 s.5/31/33/53. Most OC4IDS publication requirements are already in law.",
    "Source forward-mapping is at 97.9%, 47 of 48 ZPPA e-GP data elements map cleanly to an OC4IDS path.",
    "Procurement phase is the strongest: 20 of 153 Contracting Process fields populated (13.1%), with clean source paths like zppa_egp.contracts.contracts.id.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "+  ", bold=True, size_pt=10.5, color=LOW_TEXT)
    add_run(p, s, size_pt=10.5, color=CHARCOAL)

heading(doc, "Top three gaps", level=2)
for g in [
    "Implementation stage is effectively blank: 2% of fields populated. Contract variations, payment certificates, quality assurance reports, progress updates, none are captured in ZPPA.",
    "Maintenance and decommissioning return zero mapped fields. Beyond-completion transparency requires new capture workflows.",
    "Parties and Linked Releases sheets are silent: 18 of 968 Parties fields (1.9%) and 0 of 6 Linked Release fields mapped. Role, beneficial-ownership, and contract-release structure are absent.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "-  ", bold=True, size_pt=10.5, color=RED)
    add_run(p, g, size_pt=10.5, color=CHARCOAL)

heading(doc, "Headline recommendation", level=2)
para(doc, "Build IPPI-Zambia as an integration layer, not a new collection system. The e-GP holds procurement; the NCC holds implementation, quality, and registration; the ATI Act 2023 mandates publication of both. Integration, not data collection, is the operative task.",
     size_pt=10.5, color=CHARCOAL, space_after=8)
para(doc, "Plain-language summary. Zambia already collects most of the information OC4IDS requires. It lives in two separate government systems. The job is to connect them, and to publish what the law already says must be public.",
     size_pt=10.5, color=DMUTED, italic=True, space_after=12)

# ─── Page 3: Scope, methodology, sample ─────────────────────────────────────
doc.add_page_break()
heading(doc, "1.  Scope, methodology, and sample", level=1)

heading(doc, "1.1  Systems reviewed", level=2)
scope = doc.add_table(rows=5, cols=4)
scope.autofit = False
widths = [Cm(4.5), Cm(3), Cm(4), Cm(5.5)]
for row in scope.rows:
    for c, w in zip(row.cells, widths):
        c.width = w
# Header
hdr = ["System", "Owner", "Access method", "Coverage"]
for i, h in enumerate(hdr):
    c = scope.rows[0].cells[i]
    set_cell_shading(c, TABLE_RED)
    add_run(c.paragraphs[0], h, bold=True, size_pt=10, color=WHITE)
rows = [
    ("ZPPA e-GP System\neprocure.zppa.org.zm", "Zambia Public Procurement Authority", "Web portal + database extract", "Full procurement stage"),
    ("NCC Inspection Forms", "National Council for Construction", "Paper and digital forms", "Implementation: progress, quality, safety"),
    ("NCC Contractor Register", "National Council for Construction", "Structured database", "Party data: grade, class, performance"),
    ("NCC Project Register", "National Council for Construction", "Structured database", "Project identification and contract data"),
]
for ri, r in enumerate(rows, start=1):
    for ci, v in enumerate(r):
        c = scope.rows[ri].cells[ci]
        set_cell_shading(c, LGRAY if ri % 2 == 0 else WHITE)
        set_cell_borders(c)
        add_run(c.paragraphs[0], v, size_pt=9.5, color=CHARCOAL)

para(doc, "", space_after=8)
heading(doc, "1.2  Sample", level=2)
para(doc, "ZPPA e-GP system analysis covered all procurement records published via eprocure.zppa.org.zm for financial years 2022/23 and 2023/24, with spot-checks against live records accessed on 2026-03-14, 2026-03-18, and 2026-03-21. Sector coverage: transport, energy, water and public buildings. Sampling method: universe review of published OC4IDS-mappable fields; purposive sample of 50 live procurement records for field-population checks.",
     size_pt=10.5, space_after=8)

heading(doc, "1.3  Methodology", level=2)
method = [
    ("Standard version", "OC4IDS v0.9.5 (January 2024 release). The team did not map against earlier versions; no deprecation impact."),
    ("Evaluation criteria", "Each field classified as Populated / Partial / Missing / Not Applicable. NA entries carry a one-line justification (e.g. sector-irrelevant, or replaced by a country-specific field)."),
    ("Verification", "Mapped fields cross-checked against 50 live procurement records for field-population rates. Source system paths traced to the database schema, not inferred from portal display."),
    ("Legal alignment test", "Each OC4IDS field cross-checked against ATI Act 2023 s.8, ZPPA Act No.3/2020 s.67 and s.70, NCC Act No.10/2020 s.5 and s.53 for statutory publication authority."),
    ("Gap classification", "Overlay 8-category typology applied (no data / outside system / unstructured / poor quality / unlinkable / collected-not-disclosed / blocked / country-specific)."),
    ("Reproducibility", "Coverage numbers are re-derivable from the attached template by counting populated 'Mapping' column cells per sheet."),
]
mt = doc.add_table(rows=len(method), cols=2)
mt.autofit = False
for r in mt.rows:
    r.cells[0].width = Cm(4.2)
    r.cells[1].width = Cm(12.8)
for i, (k, v) in enumerate(method):
    kc = mt.rows[i].cells[0]
    vc = mt.rows[i].cells[1]
    for c in (kc, vc):
        set_cell_borders(c)
    add_run(kc.paragraphs[0], k, bold=True, size_pt=10, color=DMUTED)
    add_run(vc.paragraphs[0], v, size_pt=10, color=CHARCOAL)

# ─── Page 4: Quantitative backbone ──────────────────────────────────────────
doc.add_page_break()
heading(doc, "2.  Field-level coverage", level=1)

para(doc, "Zambia's template maps 73 of 1,480 OC4IDS slots (4.9%). Two sheets stand out as strong; two as silent. All mapped fields are substantive (no placeholders).",
     size_pt=10.5, space_after=8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(CHARTS / "01-sheet-coverage.png"), width=Cm(15))
para(doc, "", space_after=4)

# Per-sheet table
st = doc.add_table(rows=6, cols=5)
st.autofit = False
widths = [Cm(4.5), Cm(2.8), Cm(2.8), Cm(2.8), Cm(4.1)]
for r in st.rows:
    for c, w in zip(r.cells, widths):
        c.width = w
for i, h in enumerate(["Sheet", "Mapped", "Slots", "%", "Dominant gap type"]):
    c = st.rows[0].cells[i]
    set_cell_shading(c, TABLE_RED)
    add_run(c.paragraphs[0], h, bold=True, size_pt=10, color=WHITE)
sheet_rows = [
    ("Projects", "35", "353", "9.9%", "Preparation + post-completion absent"),
    ("Contracting Processes", "20", "153", "13.1%", "Award/implementation subpaths silent"),
    ("Parties", "18", "968", "1.9%", "Beneficial ownership, roles absent"),
    ("Linked Releases", "0", "6", "0.0%", "No release package structure in use"),
    ("Total", "73", "1,480", "4.9%", ""),
]
for ri, r in enumerate(sheet_rows, start=1):
    for ci, v in enumerate(r):
        c = st.rows[ri].cells[ci]
        set_cell_borders(c)
        bold = ri == len(sheet_rows)
        set_cell_shading(c, LGRAY if ri % 2 == 0 and not bold else (LIGHT_NEUTRAL if bold else WHITE))
        add_run(c.paragraphs[0], v, bold=bold, size_pt=10, color=CHARCOAL)

para(doc, "", space_after=10)

heading(doc, "Source forward-mapping", level=2)
para(doc, "47 of 48 source data elements in the ZPPA e-GP schema map forward to an OC4IDS path (97.9%). The single unmapped element, a local project-reference sub-field, is preserved as a country-specific extension (see §4).",
     size_pt=10.5, space_after=12)

# ─── Page 5: Phase findings ─────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "3.  Findings by lifecycle phase", level=1)

para(doc, "Each phase is scored against the OC4IDS fields nominally applicable to it. Implementation, Maintenance, and Decommissioning disclose almost nothing.",
     size_pt=10.5, space_after=6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(CHARTS / "02-phase-coverage.png"), width=Cm(15))
para(doc, "", space_after=6)

phase_rows = [
    ("Identification", "28%", "Type-1 / type-2 mix", "Reference number, owner, sector, location, purpose, description, present. Subsector and project brief, absent in system."),
    ("Preparation", "12%", "Type-1 / type-6", "Scope, budget, approval date, present. Feasibility, ESIA, funding sources, procurement plan, absent from e-GP; partial in line ministries."),
    ("Procurement", "21%", "Type-6 (mostly in-system but not disclosed)", "Tender, award, contract terms, present with clean source paths. Tender evaluation detail, contract documents, present in e-GP but not publicly disclosed."),
    ("Implementation", "2%", "Type-6 (the critical gap)", "Variation, payment, progress, quality, safety, collected by NCC under NCC Act s.53, not shared with ZPPA. Remediation is integration, not collection."),
    ("Completion", "8%", "Type-6", "Current status, projected scope, present. Final cost, completion date actuals, audit and evaluation, exist in NCC and audit bodies, not integrated."),
    ("Maintenance", "0%", "Type-1 (not collected)", "No schedule, no cost, no inspection data captured. New workflow required."),
    ("Decommissioning", "0%", "Type-1 (not collected)", "No plans, no environmental assessment, no cost data captured. New workflow required."),
]

pt = doc.add_table(rows=len(phase_rows) + 1, cols=4)
pt.autofit = False
widths = [Cm(3.0), Cm(2.0), Cm(3.8), Cm(8.2)]
for r in pt.rows:
    for c, w in zip(r.cells, widths):
        c.width = w
for i, h in enumerate(["Phase", "Coverage", "Dominant root cause", "Representative examples"]):
    c = pt.rows[0].cells[i]
    set_cell_shading(c, TABLE_RED)
    add_run(c.paragraphs[0], h, bold=True, size_pt=10, color=WHITE)
for ri, r in enumerate(phase_rows, start=1):
    for ci, v in enumerate(r):
        c = pt.rows[ri].cells[ci]
        set_cell_borders(c)
        set_cell_shading(c, LGRAY if ri % 2 == 0 else WHITE)
        if ci == 0:
            add_run(c.paragraphs[0], v, bold=True, size_pt=9.5, color=CHARCOAL)
        elif ci == 1:
            val = float(v.rstrip('%'))
            colour = RED if val < 10 else (YELLOW if val < 20 else LOW_TEXT)
            add_run(c.paragraphs[0], v, bold=True, size_pt=10, color=colour)
        else:
            add_run(c.paragraphs[0], v, size_pt=9.5, color=CHARCOAL)

# ─── Page 6: Source-field provenance sample ─────────────────────────────────
doc.add_page_break()
heading(doc, "4.  Source-field provenance (extract)", level=1)

para(doc, "Each OC4IDS path that Zambia can publish today traces to a specific ZPPA e-GP source path. The full provenance table is in Annex C; a representative sample is below.",
     size_pt=10.5, space_after=8)

prov_rows = [
    ("/id", "zppa_egp.project.project.id", "Populated", "Publish now"),
    ("/updated", "zppa_egp.project.updated", "Populated", "Publish now"),
    ("/publicAuthority/name", "zppa_egp.project.publicAuthority.name", "Populated", "Publish now"),
    ("/title", "zppa_egp.project.title", "Populated", "Publish now"),
    ("/description", "zppa_egp.project.description", "Partial", "Light transformation (free-text standardisation)"),
    ("contractingProcesses/id", "zppa_egp.contracts.contracts.id", "Populated", "Publish now"),
    ("contractingProcesses/summary/title", "zppa_egp.contracts.contracts.title", "Populated", "Publish now"),
    ("contractingProcesses/summary/tender/value", "zppa_egp.tender.tender.value", "Populated", "Publish now"),
    ("contractingProcesses/summary/award/value", "zppa_egp.contracts.contracts.value", "Populated", "Publish now"),
    ("contractingProcesses/summary/modifications", "(absent in e-GP; present in NCC Act s.53 monitoring data)", "Missing, type 6", "System change + integration"),
    ("parties/id", "zppa_egp.parties.parties.id", "Populated", "Publish now"),
    ("parties/identifier/scheme", "zppa_egp.parties.parties.identifier.scheme", "Populated", "Publish now"),
    ("parties/beneficialOwnership/*", "(absent in both systems)", "Missing, type 1 / 7", "Legal + system change"),
]

pr = doc.add_table(rows=len(prov_rows) + 1, cols=4)
pr.autofit = False
widths = [Cm(4.0), Cm(6.0), Cm(2.5), Cm(4.5)]
for r in pr.rows:
    for c, w in zip(r.cells, widths):
        c.width = w
for i, h in enumerate(["OC4IDS path", "ZPPA e-GP source path", "Status", "Disclosure pathway"]):
    c = pr.rows[0].cells[i]
    set_cell_shading(c, TABLE_RED)
    add_run(c.paragraphs[0], h, bold=True, size_pt=10, color=WHITE)
for ri, r in enumerate(prov_rows, start=1):
    for ci, v in enumerate(r):
        c = pr.rows[ri].cells[ci]
        set_cell_borders(c)
        set_cell_shading(c, LGRAY if ri % 2 == 0 else WHITE)
        sz = 9
        if ci == 0:
            add_run(c.paragraphs[0], v, size_pt=sz, color=CHARCOAL, font="Courier New")
        elif ci == 1:
            color = CHARCOAL if "absent" not in v else DMUTED
            add_run(c.paragraphs[0], v, size_pt=sz, color=color, font="Courier New", italic=("absent" in v))
        elif ci == 2:
            col = LOW_TEXT if v == "Populated" else (YELLOW if "Partial" in v else RED)
            add_run(c.paragraphs[0], v, bold=True, size_pt=sz, color=col)
        else:
            add_run(c.paragraphs[0], v, size_pt=sz, color=CHARCOAL)

para(doc, "", space_after=8)

# Country-specific note
heading(doc, "Country-specific fields retained as extensions", level=2)
para(doc, "Zambia's e-GP uses a local project-reference format that does not fit OC4IDS's id pattern. The team recommends preserving it as an extension and proposing a country-context schema note to OCP for future OC4IDS revisions. Full extension list in Annex D.",
     size_pt=10.5, space_after=12)

# ─── Page 7: Gap typology + legal crosswalk ────────────────────────────────
doc.add_page_break()
heading(doc, "5.  Why the gaps exist: typology and legal lever", level=1)

para(doc, "Each material gap is classified by root cause. Different root causes take different fixes. The two dominant categories in Zambia are type-6 (collected but not disclosed) and type-1 (not collected at all).",
     size_pt=10.5, space_after=10)

typo_rows = [
    ("Type 1", "No data exists", "Maintenance schedules, decommissioning plans.", "New capture workflow + budget line."),
    ("Type 2", "Outside reviewed system", "Feasibility studies, audit reports.", "Add to source inventory; integrate."),
    ("Type 3", "Unstructured documents only", "Project briefs, ESIA narratives.", "Fielded metadata at upload."),
    ("Type 4", "Poor quality", "Some tender evaluation free-text fields.", "Standardisation before publication."),
    ("Type 5", "Cannot be linked", "Contractor IDs across registers.", "Master data management; common identifier."),
    ("Type 6", "Collected but not disclosed", "Contract variations, payment certs, progress reports (held by NCC).", "Integration + statutory disclosure."),
    ("Type 7", "Workflow or legal constraint", "Beneficial ownership.", "Requires law or regulation change."),
    ("Type 8", "Country-specific, not in template", "Local project-reference scheme.", "Preserve as extension; propose upstream."),
]
tt = doc.add_table(rows=len(typo_rows) + 1, cols=4)
tt.autofit = False
widths = [Cm(1.5), Cm(4.5), Cm(6.0), Cm(5.0)]
for r in tt.rows:
    for c, w in zip(r.cells, widths):
        c.width = w
for i, h in enumerate(["#", "Category", "Zambia example", "Remediation"]):
    c = tt.rows[0].cells[i]
    set_cell_shading(c, TABLE_RED)
    add_run(c.paragraphs[0], h, bold=True, size_pt=10, color=WHITE)
for ri, r in enumerate(typo_rows, start=1):
    for ci, v in enumerate(r):
        c = tt.rows[ri].cells[ci]
        set_cell_borders(c)
        set_cell_shading(c, LGRAY if ri % 2 == 0 else WHITE)
        # Accent type 6 (the dominant one)
        is_t6 = r[0] == "Type 6"
        if ci == 0:
            add_run(c.paragraphs[0], v, bold=True, size_pt=9.5, color=RED if is_t6 else CHARCOAL)
        else:
            add_run(c.paragraphs[0], v, bold=is_t6 and ci == 1, size_pt=9.5, color=CHARCOAL)

para(doc, "", space_after=10)

heading(doc, "5.1  Legal anchoring", level=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(CHARTS / "04-legal-crosswalk.png"), width=Cm(15))
para(doc, "", space_after=4)

para(doc, "The Access to Information Act 2023 s.8 already mandates proactive publication of 46 OC4IDS fields. The NCC Act No.10/2020 s.53 authorises monitoring data collection that covers a further 55 fields. The ZPPA Act No.3/2020 s.67 and s.70 cover 20 procurement-transparency fields. Full crosswalk in Annex B.",
     size_pt=10.5, space_after=6)

para(doc, "Implication: publishing most of the missing implementation data does not require a new law. It requires an integration layer and a statutory instrument clarifying inter-agency data sharing.",
     size_pt=10.5, bold=True, color=CHARCOAL, space_after=12)

# ─── Page 8: Decision summary + recommendations ─────────────────────────────
doc.add_page_break()
heading(doc, "6.  Decision summary: what Zambia can publish now, and what needs to change", level=1)

para(doc, "Every publishable element is classified into one of four buckets. Government and MSG readers should be able to act on this table without reading the rest of the report.",
     size_pt=10.5, space_after=8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(CHARTS / "03-decision-buckets.png"), width=Cm(15))
para(doc, "", space_after=4)

buckets = [
    ("(a) Publish now", "Fields already mapped in e-GP with clean source paths and legal authority.",
     "35 fields. Project IDs, titles, sectors, publicAuthority, procurement method, tender and contract values, contractor IDs.",
     LOW_TEXT),
    ("(b) Publish after light transformation", "Fields present but needing format conversion, standardisation, or codelist alignment.",
     "20 fields. Free-text descriptions, date-format reconciliation, sector-code mapping to OC4IDS codelists.",
     BLUE),
    ("(c) Requires system change", "Fields NCC already collects under NCC Act s.53 that need an integration layer to reach the public.",
     "73 fields. Contract variations, progress updates, quality and safety reports, disbursement certificates, final costs.",
     YELLOW),
    ("(d) Requires legal or institutional action", "Fields blocked by absence of a statutory instrument, regulation, or workflow.",
     "25 fields. Beneficial ownership, maintenance schedules, decommissioning plans, some audit and evaluation disclosures.",
     RED),
]

bt = doc.add_table(rows=len(buckets), cols=3)
bt.autofit = False
widths = [Cm(4.5), Cm(6.5), Cm(6.0)]
for r in bt.rows:
    for c, w in zip(r.cells, widths):
        c.width = w
for ri, (label, what, example, col) in enumerate(buckets):
    cells = bt.rows[ri].cells
    set_cell_borders(cells[0], left="single", color=hex_of(col), sz="32")
    for c in cells:
        set_cell_shading(c, LGRAY)
        set_cell_margins(c, top=140, bottom=140, left=180, right=180)
    add_run(cells[0].paragraphs[0], label, bold=True, size_pt=11, color=col)
    add_run(cells[1].paragraphs[0], what, size_pt=10, color=CHARCOAL)
    add_run(cells[2].paragraphs[0], example, size_pt=10, color=CHARCOAL)

para(doc, "", space_after=12)

# ─── Page 9: Recommendations ────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "7.  Recommendations", level=1)

para(doc, "Each recommendation carries a stable identifier (R1, R2, ...), a named owner, a timeline, an effort level, and the specific legal provision it operationalises. The implementation roadmap in §8 references these IDs.",
     size_pt=10.5, space_after=10)

recs = [
    # (id, layer, title, detail, owner, legal, timeline, effort)
    ("R1", "Immediate no-regret fix",
     "Publish the 35 'publish now' fields through an OC4IDS-compliant feed on eprocure.zppa.org.zm.",
     "Use existing e-GP source paths (provenance in Annex C). No new data collection. CC BY 4.0 licence.",
     "ZPPA Director of IT", "ATI Act 2023 s.8; ZPPA Act 2020 s.67", "Q3 2026", "Low"),
    ("R2", "Immediate no-regret fix",
     "Close 20 'light transformation' fields by aligning codelists and date formats.",
     "Sector code mapping, date-format reconciliation, free-text standardisation. Script-level changes only.",
     "ZPPA Director of IT + CoST Zambia ODS lead", "ATI Act 2023 s.8", "Q3 2026", "Low"),
    ("R3", "System / configuration change",
     "Build IPPI-Zambia as the integration layer between ZPPA e-GP and NCC registers.",
     "Connects e-GP procurement records with NCC inspection, contractor, and project data. Target: 73 currently-blocked fields.",
     "NCC CEO + ZPPA Director General, CoST Zambia as delivery partner", "NCC Act 2020 s.5(k), s.5(l), s.53; ATI Act 2023 s.8",
     "Q1 2027 prototype; Q3 2027 production", "High"),
    ("R4", "System / configuration change",
     "Add variation, progress, payment, and quality-assurance capture to the NCC inspection workflow.",
     "Structured fields replacing PDF-only inspection reports. Output feeds IPPI-Zambia and public OC4IDS feed.",
     "NCC Director of Construction Monitoring", "NCC Act 2020 s.53", "Q4 2026", "Medium"),
    ("R5", "Institutional / process change",
     "Sign a statutory instrument authorising NCC to share OC4IDS-relevant data with ZPPA for publication.",
     "Closes the \"collected but not disclosed\" type-6 gap across 73 fields without new legislation.",
     "Minister of Infrastructure, Housing and Urban Development", "NCC Act 2020 s.5; ATI Act 2023 s.8", "Q4 2026", "Medium"),
    ("R6", "Institutional / process change",
     "Designate an information officer under ATI Act 2023 s.9 at each of ZPPA, NCC, and CoST Zambia.",
     "Creates an accountable publication owner in each agency. Enables cross-agency escalation.",
     "Permanent Secretary, Ministry of Finance (cross-agency coordination)", "ATI Act 2023 s.9",
     "Q3 2026", "Low"),
    ("R7", "Policy / legal change",
     "Introduce beneficial-ownership disclosure requirements via regulation under ZPPA Act No.3/2020.",
     "Closes the single largest type-7 gap (workflow / legal constraint). Covers 18+ OC4IDS Parties fields.",
     "ZPPA Director General + Attorney General's Chambers", "ZPPA Act 2020 s.67; proposed new regulation",
     "Q2 2027", "High"),
    ("R8", "Policy / legal change",
     "Amend NCC regulations to require maintenance and decommissioning data capture on all Grade-1 projects.",
     "Closes the type-1 gap on post-completion transparency (currently 0 of ~40 fields).",
     "NCC CEO + Ministry of Infrastructure", "NCC Act 2020 s.53; proposed amendment", "Q3 2027", "Medium"),
]

rt = doc.add_table(rows=len(recs) + 1, cols=7)
rt.autofit = False
widths = [Cm(0.9), Cm(2.5), Cm(4.3), Cm(2.7), Cm(2.8), Cm(1.8), Cm(1.5)]
for r in rt.rows:
    for c, w in zip(r.cells, widths):
        c.width = w
for i, h in enumerate(["ID", "Layer", "Action", "Owner", "Legal basis", "Timeline", "Effort"]):
    c = rt.rows[0].cells[i]
    set_cell_shading(c, TABLE_RED)
    add_run(c.paragraphs[0], h, bold=True, size_pt=9.5, color=WHITE)
layer_colors = {
    "Immediate no-regret fix": LOW_TEXT,
    "System / configuration change": BLUE,
    "Institutional / process change": YELLOW,
    "Policy / legal change": RED,
}
for ri, (rid, layer, title, detail, owner, legal, tl, effort) in enumerate(recs, start=1):
    cells = rt.rows[ri].cells
    set_cell_shading(cells[0], LGRAY if ri % 2 == 0 else WHITE)
    for c in cells:
        set_cell_borders(c)
        set_cell_shading(c, LGRAY if ri % 2 == 0 else WHITE)
    add_run(cells[0].paragraphs[0], rid, bold=True, size_pt=9.5, color=CHARCOAL)
    add_run(cells[1].paragraphs[0], layer, bold=True, size_pt=9, color=layer_colors[layer])
    p_action = cells[2].paragraphs[0]
    add_run(p_action, title, bold=True, size_pt=9.5, color=CHARCOAL)
    p_action2 = cells[2].add_paragraph()
    p_action2.paragraph_format.space_before = Pt(2)
    add_run(p_action2, detail, size_pt=9, color=DMUTED, italic=True)
    add_run(cells[3].paragraphs[0], owner, size_pt=9, color=CHARCOAL)
    add_run(cells[4].paragraphs[0], legal, size_pt=9, color=CHARCOAL)
    add_run(cells[5].paragraphs[0], tl, bold=True, size_pt=9.5, color=RED)
    add_run(cells[6].paragraphs[0], effort, size_pt=9, color=CHARCOAL)

para(doc, "", space_after=12)

# ─── Page 10: Roadmap ───────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "8.  Implementation roadmap", level=1)

para(doc, "The roadmap translates recommendations into a dated sequence. Every row references a recommendation ID; emergent items are labelled so. Budgets are order-of-magnitude.",
     size_pt=10.5, space_after=8)

roadmap = [
    # (priority, action, owner, dep, timeline, ref, budget)
    ("Short-term (0 to 3 months)", "Publish OC4IDS feed with 'publish now' 35 fields", "ZPPA Director of IT",
     "R1, R2", "Q3 2026", "USD 15 to 25k (integration work)"),
    ("Short-term (0 to 3 months)", "Designate ATI information officers at ZPPA, NCC, CoST Zambia",
     "Permanent Secretary, MoF", "R6", "Q3 2026", "Nil (existing staff)"),
    ("Medium-term (3 to 12 months)", "NCC inspection form redesign (structured progress + variation fields)",
     "NCC Director of Construction Monitoring", "R4", "Q4 2026", "USD 40 to 60k (form redesign + training)"),
    ("Medium-term (3 to 12 months)", "Statutory instrument on NCC-ZPPA data sharing", "Minister of Infrastructure",
     "R5", "Q4 2026", "Nil (policy)"),
    ("Medium-term (3 to 12 months)", "IPPI-Zambia prototype (OC4IDS feed from integrated e-GP + NCC data)",
     "NCC CEO + ZPPA DG, CoST Zambia delivery", "R3", "Q1 2027", "USD 150 to 220k (platform build)"),
    ("Long-term (12+ months)", "IPPI-Zambia production release",
     "NCC CEO + ZPPA DG", "R3, R4, R5", "Q3 2027", "USD 80 to 120k (hardening + hosting)"),
    ("Long-term (12+ months)", "Beneficial-ownership regulation under ZPPA Act",
     "ZPPA DG + Attorney General", "R7", "Q2 2027", "Nil (regulation)"),
    ("Long-term (12+ months)", "NCC regulation amendment on maintenance/decommissioning capture",
     "NCC CEO + Ministry of Infrastructure", "R8", "Q3 2027", "USD 20 to 40k (regulation + roll-out)"),
]

rmt = doc.add_table(rows=len(roadmap) + 1, cols=6)
rmt.autofit = False
widths = [Cm(3.0), Cm(4.5), Cm(3.2), Cm(1.3), Cm(1.6), Cm(3.2)]
for r in rmt.rows:
    for c, w in zip(r.cells, widths):
        c.width = w
for i, h in enumerate(["Priority", "Action", "Owner", "Ref", "Timeline", "Budget (USD)"]):
    c = rmt.rows[0].cells[i]
    set_cell_shading(c, TABLE_RED)
    add_run(c.paragraphs[0], h, bold=True, size_pt=9.5, color=WHITE)
for ri, r in enumerate(roadmap, start=1):
    cells = rmt.rows[ri].cells
    set_cell_shading(cells[0], LGRAY if ri % 2 == 0 else WHITE)
    for c in cells:
        set_cell_borders(c)
        set_cell_shading(c, LGRAY if ri % 2 == 0 else WHITE)
    add_run(cells[0].paragraphs[0], r[0], bold=True, size_pt=9.5, color=RED if "Short" in r[0] else (BLUE if "Medium" in r[0] else CHARCOAL))
    add_run(cells[1].paragraphs[0], r[1], size_pt=9.5, color=CHARCOAL)
    add_run(cells[2].paragraphs[0], r[2], size_pt=9, color=CHARCOAL)
    add_run(cells[3].paragraphs[0], r[3], bold=True, size_pt=9.5, color=CHARCOAL)
    add_run(cells[4].paragraphs[0], r[4], bold=True, size_pt=9.5, color=RED)
    add_run(cells[5].paragraphs[0], r[5], size_pt=9, color=CHARCOAL)

para(doc, "", space_after=12)

# ─── Page 11: Risks + review trail ──────────────────────────────────────────
doc.add_page_break()
heading(doc, "9.  Risks and mitigations", level=1)
risks = [
    ("Political will on inter-agency data sharing.", "R5 requires ministerial signature. Without it, R3 stalls.",
     "Pre-brief through CoST Zambia MSG; anchor request in ATI Act 2023 s.8 which already mandates proactive disclosure."),
    ("IPPI-Zambia scope creep.", "Tempting to build a full monitoring platform rather than an integration layer.",
     "Lock scope: OC4IDS publication + public API + download endpoint. Everything else Phase 2."),
    ("NCC form redesign resistance.", "Structured fields change inspector workflows.",
     "Co-design with NCC inspectors; pilot on 20 Grade-1 projects before full roll-out."),
    ("ATI information officer capacity.", "Newly designated officers lack OC4IDS training.",
     "Two-day CoST Zambia training workshop at designation; quarterly refresher in year 1."),
    ("Legal reform timing on beneficial ownership (R7).", "Regulatory processes can run 12 to 18 months.",
     "Start consultation in Q3 2026 so production is ready when platform launches in Q3 2027."),
]
rt = doc.add_table(rows=len(risks) + 1, cols=3)
rt.autofit = False
widths = [Cm(5.5), Cm(5.5), Cm(6.0)]
for r in rt.rows:
    for c, w in zip(r.cells, widths):
        c.width = w
for i, h in enumerate(["Risk", "Why it matters", "Mitigation"]):
    c = rt.rows[0].cells[i]
    set_cell_shading(c, TABLE_RED)
    add_run(c.paragraphs[0], h, bold=True, size_pt=10, color=WHITE)
for ri, r in enumerate(risks, start=1):
    cells = rt.rows[ri].cells
    for c in cells:
        set_cell_borders(c)
        set_cell_shading(c, LGRAY if ri % 2 == 0 else WHITE)
    add_run(cells[0].paragraphs[0], r[0], bold=True, size_pt=9.5, color=CHARCOAL)
    add_run(cells[1].paragraphs[0], r[1], size_pt=9.5, color=CHARCOAL)
    add_run(cells[2].paragraphs[0], r[2], size_pt=9.5, color=CHARCOAL)

para(doc, "", space_after=12)

heading(doc, "10.  Review trail", level=1)
trail_rows = [
    ("Zambia Technical Team (NCC + ZPPA + CoST Zambia)", "Primary author", "2026-03-03"),
    ("CoST International Secretariat, Open Data Specialist", "Field-level validation", "2026-04-23 (Overlay gate + full rubric applied)"),
    ("CoST Zambia MSG", "Multi-stakeholder review", "Scheduled 2026-05-06"),
    ("ZPPA Director of IT", "Technical feasibility sign-off on R1, R2, R3", "Pending"),
    ("NCC Director of Construction Monitoring", "Technical feasibility sign-off on R3, R4, R8", "Pending"),
]
trt = doc.add_table(rows=len(trail_rows) + 1, cols=3)
trt.autofit = False
widths = [Cm(7.0), Cm(6.5), Cm(3.5)]
for r in trt.rows:
    for c, w in zip(r.cells, widths):
        c.width = w
for i, h in enumerate(["Reviewer", "Role", "Date / status"]):
    c = trt.rows[0].cells[i]
    set_cell_shading(c, TABLE_RED)
    add_run(c.paragraphs[0], h, bold=True, size_pt=10, color=WHITE)
for ri, r in enumerate(trail_rows, start=1):
    cells = trt.rows[ri].cells
    for c in cells:
        set_cell_borders(c)
        set_cell_shading(c, LGRAY if ri % 2 == 0 else WHITE)
    for ci, v in enumerate(r):
        add_run(cells[ci].paragraphs[0], v, size_pt=9.5, color=CHARCOAL,
                bold=(ci == 2 and "Pending" in v))

para(doc, "", space_after=14)

# ─── Page 12: Conclusion + annex references ─────────────────────────────────
heading(doc, "11.  Conclusion", level=1)
para(doc, "Zambia's infrastructure data system is not as thin as it first appears. The ZPPA e-GP holds the procurement backbone. The NCC holds the implementation record. The ATI Act 2023 already requires publication of most of both. The binding constraint is integration and disclosure, not collection.",
     size_pt=10.5, space_after=8)
para(doc, "This analysis estimates that 35 OC4IDS fields can move to public disclosure in the current quarter. A further 20 can follow with light transformation. Another 73 unlock through IPPI-Zambia, the proposed integration layer. The remaining 25 require regulation or workflow change.",
     size_pt=10.5, space_after=8)
para(doc, "The immediate path forward has three moves: publish what is already publishable, sign the NCC-ZPPA data-sharing instrument, and scope IPPI-Zambia as an integration platform, not a monitoring system. Do those three and 128 of 153 OC4IDS fields become public within twelve months.",
     size_pt=10.5, bold=True, color=CHARCOAL, space_after=14)

heading(doc, "Annexes (references)", level=2)
for a in [
    "Annex A ,  Completed OC4IDS v0.9.5 Field-Level Mapping Template (attached: 2025_11_Zambia_OC4IDS_0.9.5_Field-Level_Mapping_Template.xlsx).",
    "Annex B ,  Legal crosswalk: ATI Act 2023 / ZPPA Act 2020 / NCC Act 2020 mapped to each OC4IDS field.",
    "Annex C ,  Full source-field provenance table: every mapped OC4IDS path × ZPPA e-GP source path.",
    "Annex D ,  Country-specific fields retained as extensions.",
    "Annex E ,  Sampled records used for field-population verification (n=50).",
    "Annex F ,  NCC Inspection Form data dictionary (proposed structured redesign).",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "·  ", color=MUTED, size_pt=10.5)
    add_run(p, a, size_pt=10, color=CHARCOAL)

para(doc, "", space_after=14)

# Sign-off stripe
para(doc, "CoST, the Infrastructure Transparency Initiative  ·  www.infrastructuretransparency.org  ·  Published under CC BY 4.0",
     size_pt=8, italic=True, color=LMUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

doc.save(str(OUT))
print(f"OK {OUT}")
