"""Build Doc 3: Sample final report for Zambia (Option A).

A fully-worked model report the Zambia Technical Team can adapt directly.
Every stakeholder name is a placeholder pending MSG consultation. The
structure, numbers, legal analysis, and recommendation logic are at the
target standard a CoST IS reviewer would accept.

Layout: Portrait A4, Full Publication intensity, Activity Report archetype.

Draws on:
- Zambia OC4IDS mapping template v0.9.5 (real mapped paths and counts).
- Zambia OC4IDS mapping report draft of 3 March 2026 (real legal analysis,
  NCC insight, lifecycle framework).
- The Overlay rubric's 12-section structure (C0 to C11).
"""
import pathlib
from docx import Document
from docx.shared import Cm, Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import _docx_helpers as H
from visuals.infographics import cover_photo_strip

CHARTS = pathlib.Path(__file__).parent.parent / "charts"
OUT = pathlib.Path(__file__).parent.parent / "02-sample-final-report.docx"

CONTENT_W = 16.6

# Cover photo strip composed as one PNG at exact content width.
_COVER_STRIP = CHARTS / "cover-strip-sample-final-report.png"
cover_photo_strip(
    [
        H.ASSETS / "stock-africa-highway-aerial.jpeg",
        H.ASSETS / "stock-africa-rural-road.jpeg",
        H.ASSETS / "stock-africa-scaffolding.jpeg",
    ],
    _COVER_STRIP,
    target_width_cm=CONTENT_W,
    strip_height_mm=50,
)


# ─── Document setup ─────────────────────────────────────────────────────────
doc = Document()
H.setup_portrait(doc.sections[0], margin_mm=22)
H.setup_base_style(doc)
H.add_logo_header(doc.sections[0])
H.add_stripe_footer(doc.sections[0])


# ─── Cover ──────────────────────────────────────────────────────────────────
# Single 2-row table: photo strip (row 1) + red title block (row 2). Rows share
# the table's width exactly — pixel-perfect edge alignment by construction.
cover = doc.add_table(rows=2, cols=1)
H.set_table_full_width(cover, CONTENT_W)
H.set_table_zero_cell_margins(cover)
H.apply_col_widths(cover, [CONTENT_W])

# Row 1 — photo strip
scell = cover.rows[0].cells[0]
H.set_cell_margins(scell, top=0, bottom=0, left=0, right=0)
H.set_cell_borders(scell)
scp = scell.paragraphs[0]
scp.paragraph_format.space_before = Pt(0)
scp.paragraph_format.space_after = Pt(0)
scp.paragraph_format.left_indent = Cm(0)
scp.paragraph_format.right_indent = Cm(0)
scp.paragraph_format.first_line_indent = Cm(0)
scp.alignment = WD_ALIGN_PARAGRAPH.LEFT
scp.add_run().add_picture(str(_COVER_STRIP), width=Cm(CONTENT_W))

# Row 2 — red title block
rcell = cover.rows[1].cells[0]
H.set_cell_shading(rcell, H.COVER_RED)
H.set_cell_margins(rcell, top=520, bottom=520, left=480, right=480)
H.set_cell_borders(rcell)
rcp = rcell.paragraphs[0]
rcp.paragraph_format.space_after = Pt(4)
H.add_run(rcp, "SAMPLE STRUCTURE FOR A FIELD-LEVEL MAPPING REPORT", bold=True, size_pt=12, color=H.WHITE)
rcp2 = rcell.add_paragraph()
rcp2.paragraph_format.space_after = Pt(8)
H.add_run(rcp2, "What data Zambia has, where it is held, what is missing, and what to do next",
          bold=True, size_pt=19, color=H.WHITE)
rcp3 = rcell.add_paragraph()
rcp3.paragraph_format.space_after = Pt(0)
H.add_run(rcp3, "Prepared by the Zambia Technical Team.  With support from CoST International Secretariat.",
          size_pt=10, color=H.WHITE)
rcp4 = rcell.add_paragraph()
rcp4.paragraph_format.space_before = Pt(4)
rcp4.paragraph_format.space_after = Pt(0)
H.add_run(rcp4, "Version: Model report.  Date: [to be set by Zambia team].",
          size_pt=10, color=H.WHITE, italic=True)

H.para(doc, "", space_after=10)

# Model disclaimer
H.callout_box(
    doc, CONTENT_W,
    label="HOW TO USE THIS DOCUMENT",
    label_color=H.DMUTED,
    body="This is a sample structure for a field-level mapping report. It is not a published country report. CoST IS prepared it as a drafting reference for member-country teams, using Zambia data to make the pattern concrete. Another country adopting this model should keep the structure and replace the country-specific content (statutes, institutions, systems, stakeholder names) with its own. Country-specific names shown in square brackets are placeholders pending MSG consultation. Any legal interpretation or indicative figure in this sample requires country validation before use.",
)
H.para(doc, "", space_after=10)

# What this report answers — the four practical questions
H.callout_box(
    doc, CONTENT_W,
    label="WHAT THIS REPORT ANSWERS",
    label_color=H.RED,
    body="A field-level mapping report exists to help a country answer four practical questions: What data do we already have? Where is it held? What is missing? What do we do next? This report is organised around those four questions. Legal analysis, platform design, budget planning, and reform sequencing are deliberately not carried inside this mapping report so the core mapping story stays clear; countries that want that strategic layer should commission it separately.",
)
H.para(doc, "", space_after=10)

# Publisher/metadata block
meta = doc.add_table(rows=5, cols=2)
H.set_table_full_width(meta, CONTENT_W)
H.apply_col_widths(meta, [4.8, CONTENT_W - 4.8])
for i, (k, v) in enumerate([
    ("Publisher:", "National Council for Construction (NCC) and Zambia Public Procurement Authority (ZPPA); coordinated by CoST Zambia, hosted by NCC."),
    ("Prepared by:", "Zambia Technical Team. [Lead author: to be confirmed by MSG.] With support from CoST International Secretariat."),
    ("Standards reference:", "OC4IDS v0.9.5 field-level mapping template. CoST Infrastructure Data Standard (IDS)."),
    ("Systems reviewed:", "ZPPA e-GP System (eprocure.zppa.org.zm), NCC Inspection Forms, NCC Contractor Register, NCC Project Register."),
    ("Publication licence:", "CC BY 4.0. Published at [portal URL to be set at IPPI-Zambia launch]."),
]):
    kc = meta.rows[i].cells[0]
    vc = meta.rows[i].cells[1]
    for c in (kc, vc):
        H.set_cell_borders(c)
        H.set_cell_margins(c, top=140, bottom=140, left=180, right=180)
    H.add_run(kc.paragraphs[0], k, bold=True, size_pt=9.5, color=H.DMUTED)
    H.add_run(vc.paragraphs[0], v, size_pt=9.5, color=H.CHARCOAL)

H.para(doc, "CoST, the Infrastructure Transparency Initiative  •  Company number 8159144  •  Charity number 1152236",
       size_pt=7.5, italic=True, color=H.LMUTED,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)


# ─── Page 2: Executive summary ──────────────────────────────────────────────
doc.add_page_break()
H.heading(doc, "Executive summary", level=1)

H.verdict_box(
    doc, CONTENT_W,
    label="VERDICT",
    headline="Zambia's public procurement system can publish 35 OC4IDS fields today. Another 20 will follow after small format fixes. Most of the information that is currently missing (contract variations, progress reports, payment records, quality assurance) is already collected by the National Council for Construction under the NCC Act 2020, section 53.",
    tail="On the evidence of this mapping, the problem is better read as a disclosure gap than as a data gap. The Access to Information Act 2023 section 8 appears to mandate proactive publication of most of the missing fields, subject to final legal review. The three immediate actions are set out in section 8; scoping the integration architecture that would carry them is a separate product.",
)
H.para(doc, "", space_after=12)

H.heading(doc, "Headline numbers", level=2)

stats = doc.add_table(rows=1, cols=3)
H.set_table_full_width(stats, CONTENT_W)
col = CONTENT_W / 3
H.apply_col_widths(stats, [col, col, col])
H.set_row_height(stats.rows[0], 4.8)

for i, (val, label, col_color) in enumerate([
    ("35",
     "OC4IDS fields the ZPPA e-GP system can publish today, with no new collection work and no system changes.",
     H.LOW_TEXT),
    ("55",
     "OC4IDS fields publishable within 12 months (35 today plus 20 after format fixes).",
     H.BLUE),
    ("46",
     "OC4IDS fields the Access to Information Act 2023 section 8 already mandates for proactive disclosure.",
     H.YELLOW),
]):
    cell = stats.rows[0].cells[i]
    H.set_cell_shading(cell, H.LGRAY)
    H.set_cell_margins(cell, top=280, bottom=280, left=280, right=280)
    H.set_cell_borders(cell, top="single", color=H.hex_of(col_color), sz="40")
    p1 = cell.paragraphs[0]
    H.add_run(p1, val, bold=True, size_pt=26, color=col_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    H.add_run(p2, label, size_pt=9, color=H.DMUTED)

H.para(doc, "", space_after=12)

H.heading(doc, "Top three strengths", level=2)
for s in [
    "Procurement data coverage is consistent. The ZPPA e-GP system captures 20 of 153 Contracting Process fields (13.1%) with clean source paths that can be reused directly for an OC4IDS feed. All mapped fields are substantive (no placeholders).",
    "The legal mandate is already in place. The Access to Information Act 2023 section 8 requires proactive publication of public contracts, suppliers, and amounts. The NCC Act 2020 section 53 authorises the National Council for Construction to monitor projects and share records. Most of what OC4IDS asks for is what the law already requires.",
    "Source-field traceability is excellent. 47 of 48 source data elements in the ZPPA e-GP schema map forward to an OC4IDS path (97.9%). Every mapped field can be verified back to a specific database path.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    H.add_run(p, "+  ", bold=True, size_pt=10.5, color=H.LOW_TEXT)
    H.add_run(p, s, size_pt=10.5, color=H.CHARCOAL)

H.heading(doc, "Top three gaps", level=2)
for g in [
    "The implementation phase is nearly blank. Only 2% of Implementation-phase OC4IDS fields are populated. Contract variations, payment certificates, progress updates, and quality assurance reports are absent from the e-GP system.",
    "Two OC4IDS sheets are silent. The Linked Releases sheet has 0 of 6 fields populated; the Parties sheet has 18 of 968. Beneficial ownership, party roles, and contract-release structure are not captured.",
    "Maintenance and decommissioning return zero mapped fields. Post-completion transparency requires new capture workflows; these do not currently exist in any reviewed system.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    H.add_run(p, "-  ", bold=True, size_pt=10.5, color=H.RED)
    H.add_run(p, g, size_pt=10.5, color=H.CHARCOAL)

H.heading(doc, "Headline reading", level=2)
H.para(doc, "Most of the missing implementation data already exists inside government: the ZPPA e-GP holds the procurement record, and the NCC holds the implementation record. The constraint, on the evidence reviewed, is disclosure and integration rather than collection. Three directional moves follow; their detailed scoping, owners, and sequencing sit outside this mapping report and should be developed separately through the MSG.",
       size_pt=10.5, color=H.CHARCOAL, space_after=10)

H.para(doc, "Plain-language summary. Zambia already collects most of what OC4IDS requires. The data is held in two government systems that do not yet talk to each other. The mapping report names what exists and what is missing; the reform work of connecting them is a separate product.",
       size_pt=10.5, italic=True, color=H.DMUTED, space_after=12)


# ─── Page 3: Scope, methodology, systems ───────────────────────────────────
doc.add_page_break()
H.heading(doc, "1.  Scope, methodology, and systems reviewed", level=1)

H.heading(doc, "1.1  Scope", level=2)
H.para(doc, "This assessment evaluates Zambia's public infrastructure data systems against the Open Contracting for Infrastructure Data Standard (OC4IDS) v0.9.5. It covers:",
       size_pt=10.5, space_after=6)
for line in [
    "Procurement records in the ZPPA electronic Government Procurement (e-GP) system for financial years 2022/23 and 2023/24.",
    "Project monitoring data collected by the National Council for Construction under the NCC Act 2020.",
    "Contractor and project registration data held in NCC registers.",
    "The legal and institutional framework authorising public infrastructure disclosure.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    H.add_run(p, "•   " + line, size_pt=10.5, color=H.CHARCOAL)

H.para(doc, "Sectors covered: transport, energy, water and sanitation, and public buildings.",
       size_pt=10.5, space_after=10, space_before=4)

H.heading(doc, "1.2  Sample", level=2)
H.para(doc, "All procurement records published via eprocure.zppa.org.zm for financial years 2022/23 and 2023/24 were analysed, with spot-checks against live records accessed on 14, 18, and 21 March 2026. Sampling method: universe review of published OC4IDS-mappable fields; purposive sample of 50 live procurement records for field-population checks.\n\n[Note to Zambia team: before publication, confirm the actual record count, sector mix, and value totals for this financial-year window. If the universe figure differs materially from what the live system shows on the submission date, restate here.]",
       size_pt=10.5, space_after=10)

H.heading(doc, "1.3  Systems reviewed", level=2)
sc = doc.add_table(rows=5, cols=5)
H.set_table_full_width(sc, CONTENT_W)
H.apply_col_widths(sc, [4.0, 3.2, 3.6, 2.6, CONTENT_W - 13.4])
H.apply_default_padding(sc)
H.styled_table_header(sc.rows[0],
                      ["System", "Owner", "Access method", "Public?", "Coverage"])

for ri, r in enumerate([
    ("ZPPA e-GP System (eprocure.zppa.org.zm)",
     "Zambia Public Procurement Authority",
     "Web portal plus structured database extract",
     "Yes (portal)",
     "Full procurement stage: tender, award, contract."),
    ("NCC Inspection Forms",
     "National Council for Construction",
     "Paper and digital forms; quarterly submission",
     "No (internal)",
     "Implementation: progress, quality, safety, compliance."),
    ("NCC Contractor Register",
     "National Council for Construction",
     "Structured database",
     "Partial",
     "Party data: registered firms, grade, class, performance history."),
    ("NCC Project Register",
     "National Council for Construction",
     "Structured database",
     "Partial",
     "Project identification: reference number, owner, sector, location, contract data."),
], start=1):
    cells = sc.rows[ri].cells
    H.set_row_height(sc.rows[ri], 1.6)
    bg = H.LGRAY if ri % 2 == 0 else H.WHITE
    for c in cells:
        H.set_cell_shading(c, bg)
        H.set_cell_borders(c)
    for ci, v in enumerate(r):
        H.add_run(cells[ci].paragraphs[0], v, size_pt=9.5,
                  color=H.CHARCOAL, bold=(ci == 0))

H.para(doc, "", space_after=10)

H.heading(doc, "1.4  Methodology", level=2)
for label, body in [
    ("Standard version",
     "OC4IDS v0.9.5 (January 2024 release). The team did not map against earlier versions."),
    ("Evaluation criteria",
     "Each field classified as Populated / Partial / Missing / Not applicable. NA entries carry a one-line justification (typically sector-irrelevant, or replaced by a country-specific field)."),
    ("Verification",
     "Mapped fields were cross-checked against 50 live procurement records for field-population rates. Source system paths were traced to the database schema, not inferred from portal display."),
    ("Legal alignment test",
     "Each OC4IDS field was cross-checked against the Access to Information Act 2023 section 8, the ZPPA Act No.3/2020 sections 67 and 70, and the NCC Act No.10/2020 sections 5 and 53 for statutory publication authority."),
    ("Gap classification",
     "Gaps classified by root cause: not collected, outside reviewed system, unstructured documents only, poor quality, cannot be linked, collected but not publicly disclosed, blocked by workflow or legal constraint, or country-specific and not in the template."),
    ("Reproducibility",
     "Coverage numbers are re-derivable from the attached template by counting populated Mapping-column cells per sheet. A second reviewer running the same template against the same systems should reach the same results."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    H.add_run(p, label + ".  ", bold=True, size_pt=10.5, color=H.CHARCOAL)
    H.add_run(p, body, size_pt=10.5, color=H.CHARCOAL)


# ─── Page 4: Coverage chassis ──────────────────────────────────────────────
doc.add_page_break()
H.heading(doc, "2.  Field-level coverage", level=1)

H.para(doc, "Zambia's mapping maps 73 OC4IDS fields from the ZPPA e-GP system into the four OC4IDS sheets. Two sheets stand out as strong; two are effectively silent. All mapped fields are substantive, with no placeholders.",
       size_pt=10.5, space_after=8)

H.centered_image(doc, CHARTS / "01-sheet-coverage.png", width_cm=CONTENT_W - 1,
                 space_before=4, space_after=10)

st = doc.add_table(rows=6, cols=4)
H.set_table_full_width(st, CONTENT_W)
H.apply_col_widths(st, [4.2, 3.4, 2.8, CONTENT_W - 10.4])
H.apply_default_padding(st)
H.styled_table_header(st.rows[0],
                      ["Sheet", "Fields mapped", "Coverage", "Dominant gap type"])

for ri, r in enumerate([
    ("Projects", "35 of 353", "9.9%", "Preparation and post-completion fields absent."),
    ("Contracting Processes", "20 of 153", "13.1%", "Award and implementation subpaths silent."),
    ("Parties", "18 of 968", "1.9%", "Beneficial ownership and role data absent."),
    ("Linked Releases", "0 of 6", "0.0%", "No release-package structure in use."),
    ("Total", "73 of 1,480 template slots", "4.9%", "Applicable-scope coverage is higher (see below)."),
], start=1):
    cells = st.rows[ri].cells
    H.set_row_height(st.rows[ri], 1.4)
    is_total = ri == 5
    bg = H.LIGHT_NEUTRAL if is_total else (H.LGRAY if ri % 2 == 0 else H.WHITE)
    for c in cells:
        H.set_cell_shading(c, bg)
        H.set_cell_borders(c)
    for ci, v in enumerate(r):
        H.add_run(cells[ci].paragraphs[0], v, size_pt=10, color=H.CHARCOAL, bold=is_total)

H.para(doc, "", space_after=10)

H.callout_box(
    doc, CONTENT_W,
    label="ON THE DENOMINATORS",
    label_color=H.BLUE,
    body="The OC4IDS template contains approximately 1,480 field slots in v0.9.5 across required, recommended, and sector-specific fields; the total varies across OC4IDS versions. No publisher is expected to fill every slot; OC4IDS is a universe, not an obligation. The denominators this report uses are (a) the required fields the publisher is accountable for (all populated), (b) the subset of required and recommended fields applicable to ZPPA's scope (approximately 55 of 150 populated), and (c) the count of fields publishable in practice (35 ready to publish today, 55 within 12 months). The full OC4IDS universe total is reported in Annex A for cross-country benchmarking only.",
)
H.para(doc, "", space_after=12)

H.heading(doc, "Source forward-mapping", level=2)
H.para(doc, "47 of 48 source data elements in the ZPPA e-GP schema map forward to an OC4IDS path (97.9%). The single unmapped element is a local project-reference sub-field preserved as a country-specific extension (see section 5).",
       size_pt=10.5, space_after=12)


# ─── Page 5: Findings by phase ─────────────────────────────────────────────
doc.add_page_break()
H.heading(doc, "3.  Findings by lifecycle phase", level=1)

H.para(doc, "Disclosure is strong at identification and procurement. It thins sharply across implementation and collapses at maintenance and decommissioning.",
       size_pt=10.5, space_after=6)

H.centered_image(doc, CHARTS / "02-phase-coverage.png", width_cm=CONTENT_W - 1,
                 space_before=4, space_after=10)

phases = [
    ("Identification", "7 of ~25  (28%)",
     "Not collected / Unstructured (mixed)",
     "Project reference number, owner, sector, location, purpose, and description are captured. Sub-sector and project-brief fields are absent from e-GP and held only in line-ministry PDFs."),
    ("Preparation", "3 of ~25  (12%)",
     "Outside reviewed system",
     "Project scope, budget, and approval date are present. Feasibility studies, ESIAs, funding sources, and procurement plans are held by line ministries, not by ZPPA or NCC."),
    ("Procurement", "6 of ~28  (21%)",
     "Collected but not yet fully disclosed",
     "Procuring entity, procurement method, number of bidders, cost estimate, contract price, and contract dates are published. Tender evaluation detail, final selection justification, and contract documents sit inside e-GP but are not public."),
    ("Implementation", "1 of ~45  (~2%)",
     "Collected but not yet disclosed (NCC-held)",
     "Contract variations, progress updates, payment certificates, quality assurance, and safety data are all captured by NCC under section 53 of the NCC Act. None reach the ZPPA e-GP or the public."),
    ("Completion", "2 of ~25  (8%)",
     "Collected but not yet disclosed",
     "Project status and projected scope are published. Actual completion cost, actual completion date, reasons for deviations, and audit and evaluation reports exist in NCC and audit bodies but are not integrated into any public system."),
    ("Maintenance", "0 of ~10  (0%)",
     "Not collected",
     "Maintenance schedules, cost logs, and inspection data are not captured in any reviewed system. A new capture workflow is required."),
    ("Decommissioning", "0 of ~10  (0%)",
     "Not collected",
     "Decommissioning plans, environmental assessments, and costs are not captured. A new workflow and a regulatory requirement are both needed."),
]

pt = doc.add_table(rows=len(phases) + 1, cols=4)
H.set_table_full_width(pt, CONTENT_W)
H.apply_col_widths(pt, [2.8, 1.8, 4.2, CONTENT_W - 8.8])
H.apply_default_padding(pt)
H.styled_table_header(pt.rows[0],
                      ["Phase", "Coverage", "Dominant root cause",
                       "Representative evidence"])

import re as _re
for ri, (phase, pct_s, cause, ev) in enumerate(phases, start=1):
    cells = pt.rows[ri].cells
    H.set_row_height(pt.rows[ri], 1.6)
    bg = H.LGRAY if ri % 2 == 0 else H.WHITE
    for c in cells:
        H.set_cell_shading(c, bg)
        H.set_cell_borders(c)
    H.add_run(cells[0].paragraphs[0], phase, bold=True, size_pt=10, color=H.CHARCOAL)
    # Extract percentage for colour selection, regardless of surrounding fraction.
    m = _re.search(r"(\d+)%", pct_s)
    pct = float(m.group(1)) if m else 0.0
    col = H.RED if pct < 10 else (H.YELLOW if pct < 20 else H.LOW_TEXT)
    H.add_run(cells[1].paragraphs[0], pct_s, bold=True, size_pt=10, color=col)
    H.add_run(cells[2].paragraphs[0], cause, size_pt=9.5, color=H.CHARCOAL)
    H.add_run(cells[3].paragraphs[0], ev, size_pt=9.5, color=H.CHARCOAL)

H.para(doc, "Fractions show fields mapped against applicable slots per phase (approximately 25 to 45 slots per phase in OC4IDS v0.9.5 for ZPPA-scope reporting). Denominators are approximate; the full mapping template is in Annex A. The phase table reports only fields that map to a single lifecycle phase. Parties sheet fields (18 mapped), Linked Releases, cross-cutting metadata, and codelist fields count toward the 73 total in section 2 but are not phase-specific, so the phase numerators do not sum to 73.",
       size_pt=9, italic=True, color=H.DMUTED, space_after=10)

H.callout_box(
    doc, CONTENT_W,
    label="THE STRUCTURAL READING",
    label_color=H.RED,
    body="Four of the seven phases have 'collected but not yet disclosed' as their dominant root cause. The data exists inside government. It does not reach the public. This reframes the rest of the report: the job is not to collect more data, it is to integrate what is already collected and to publish what the law already mandates.",
)
H.para(doc, "", space_after=12)


# ─── Page 6: Where each field comes from ───────────────────────────────────
doc.add_page_break()
H.heading(doc, "4.  Where each field comes from", level=1)

H.para(doc, "Each OC4IDS path Zambia can publish traces to a specific ZPPA e-GP source path. The table below shows a representative sample. The complete source-paths list appears in Annex C.",
       size_pt=10.5, space_after=10)

prov_rows = [
    ("/id", "zppa_egp.project.project.id", "Populated", "Ready to publish"),
    ("/title", "zppa_egp.project.title", "Populated", "Ready to publish"),
    ("/description", "zppa_egp.project.description", "Partial", "Small format fix (free-text standardisation)"),
    ("/publicAuthority/name", "zppa_egp.project.publicAuthority.name", "Populated", "Ready to publish"),
    ("/sector", "zppa_egp.project.sector", "Populated", "Small format fix (codelist alignment)"),
    ("contractingProcesses/id", "zppa_egp.contracts.contracts.id", "Populated", "Ready to publish"),
    ("contractingProcesses/summary/title", "zppa_egp.contracts.contracts.title", "Populated", "Ready to publish"),
    ("contractingProcesses/summary/tender/value", "zppa_egp.tender.tender.value", "Populated", "Ready to publish"),
    ("contractingProcesses/summary/award/value", "zppa_egp.contracts.contracts.value", "Populated", "Ready to publish"),
    ("contractingProcesses/summary/modifications", "(NCC Inspection Forms, s.53 monitoring data)", "Missing from e-GP", "Needs system work (integration)"),
    ("parties/id", "zppa_egp.parties.parties.id", "Populated", "Ready to publish"),
    ("parties/identifier/scheme", "zppa_egp.parties.parties.identifier.scheme", "Populated", "Ready to publish"),
    ("parties/beneficialOwnership/*", "(not in any reviewed system)", "Not collected", "Needs policy action"),
]

pr = doc.add_table(rows=len(prov_rows) + 1, cols=4)
H.set_table_full_width(pr, CONTENT_W)
H.apply_col_widths(pr, [4.0, 5.4, 2.4, CONTENT_W - 11.8])
H.apply_default_padding(pr)
H.styled_table_header(pr.rows[0],
                      ["OC4IDS path", "ZPPA e-GP source path", "Status", "Disclosure pathway"])

for ri, r in enumerate(prov_rows, start=1):
    cells = pr.rows[ri].cells
    H.set_row_height(pr.rows[ri], 1.1)
    bg = H.LGRAY if ri % 2 == 0 else H.WHITE
    for c in cells:
        H.set_cell_shading(c, bg)
        H.set_cell_borders(c)
    # OC4IDS path in monospace
    H.add_run(cells[0].paragraphs[0], r[0], size_pt=9, color=H.CHARCOAL, font="Courier New")
    # Source path: italic if absent
    italic = "not in" in r[1] or "Inspection" in r[1]
    H.add_run(cells[1].paragraphs[0], r[1], size_pt=9,
              color=(H.DMUTED if italic else H.CHARCOAL),
              font="Courier New", italic=italic)
    # Status with colour
    status = r[2]
    if "Populated" in status:
        col = H.LOW_TEXT
    elif "Partial" in status:
        col = H.YELLOW
    else:
        col = H.RED
    H.add_run(cells[2].paragraphs[0], status, bold=True, size_pt=9.5, color=col)
    H.add_run(cells[3].paragraphs[0], r[3], size_pt=9.5, color=H.CHARCOAL)

H.para(doc, "", space_after=12)


# ─── Page 7: Country-specific fields ───────────────────────────────────────
doc.add_page_break()
H.heading(doc, "5.  Country-specific fields and extensions", level=1)

H.para(doc, "The ZPPA e-GP system captures some data points that do not map neatly to any OC4IDS path. Rather than discarding these, the report preserves them as country-specific extensions in the published feed. Three are noted here; the full list appears in Annex D.",
       size_pt=10.5, space_after=10)

ext = doc.add_table(rows=4, cols=3)
H.set_table_full_width(ext, CONTENT_W)
H.apply_col_widths(ext, [5.0, 5.5, CONTENT_W - 10.5])
H.apply_default_padding(ext)
H.styled_table_header(ext.rows[0],
                      ["Zambia-specific field",
                       "What it captures",
                       "Recommended treatment"])

for ri, r in enumerate([
    ("Local project-reference scheme",
     "A 14-character code combining sector, ministry, and year of approval, used as an internal cross-reference by the Ministry of Finance.",
     "Preserve as extension field (zambia:projectReference). Propose adding as recommended to OC4IDS schema for sector-specific identifiers."),
    ("NCC contractor grade",
     "A capability grade (1 to 6) assigned by NCC to registered contractors based on financial capacity, equipment, and past performance.",
     "Preserve as extension field on the Parties sheet (zambia:contractorGrade). Document the grading criteria in Annex D."),
    ("Force-account flag",
     "A flag indicating whether implementation is by direct government labour rather than an external contractor.",
     "Preserve as extension field on Contracting Processes (zambia:forceAccount). Propose to OCP for inclusion in OC4IDS procurement-method codelist."),
], start=1):
    cells = ext.rows[ri].cells
    H.set_row_height(ext.rows[ri], 1.8)
    bg = H.LGRAY if ri % 2 == 0 else H.WHITE
    for c in cells:
        H.set_cell_shading(c, bg)
        H.set_cell_borders(c)
    H.add_run(cells[0].paragraphs[0], r[0], bold=True, size_pt=10, color=H.CHARCOAL)
    H.add_run(cells[1].paragraphs[0], r[1], size_pt=10, color=H.CHARCOAL)
    H.add_run(cells[2].paragraphs[0], r[2], size_pt=10, color=H.CHARCOAL)

H.para(doc, "", space_after=12)


# ─── Page 8: Why the gaps exist ────────────────────────────────────────────
H.heading(doc, "6.  Why the gaps exist", level=1)

H.para(doc, "Every material gap is classified by root cause. Different causes require different fixes. The dominant cause in Zambia (four of seven lifecycle phases) is 'collected but not yet disclosed': the data exists inside government but is not public. The remediation is integration, not new collection.",
       size_pt=10.5, space_after=10)

typology = [
    ("Not collected",
     "Data does not exist in any reviewed system.",
     "Maintenance schedules; decommissioning plans.",
     "New capture workflow plus budget line."),
    ("Collected but not yet disclosed",
     "Data exists inside government, often with legal disclosure authority, but does not reach the public.",
     "Contract variations, progress reports, quality assurance, and payment certificates held by NCC.",
     "Integration between systems. A short statutory instrument may be sufficient. Dominant cause across four lifecycle phases."),
    ("Only in unstructured documents",
     "Data exists but sits inside PDFs or scanned forms, unfielded.",
     "Project briefs, feasibility studies, ESIA narratives.",
     "Fielded metadata at upload; structured replacement for inspection PDFs."),
    ("Collected inconsistently",
     "Data captured on some projects, not others; format varies.",
     "Tender evaluation free-text fields; sub-sector classifications.",
     "Standardisation before publication; codelist alignment."),
    ("Restricted by law or workflow",
     "Data exists but a specific legal provision or internal policy blocks disclosure.",
     "Beneficial ownership of contracting firms.",
     "Regulation or policy change, not a technical fix."),
]

tt = doc.add_table(rows=len(typology) + 1, cols=4)
H.set_table_full_width(tt, CONTENT_W)
col_w = (CONTENT_W - 3.6) / 3
H.apply_col_widths(tt, [3.6, col_w, col_w, col_w])
H.apply_default_padding(tt)
H.styled_table_header(tt.rows[0],
                      ["Cause", "What it means", "Zambia example", "What fixes it"])

for ri, r in enumerate(typology, start=1):
    cells = tt.rows[ri].cells
    H.set_row_height(tt.rows[ri], 1.9)
    accent = ri == 2
    bg = H.LGRAY if ri % 2 == 0 else H.WHITE
    for c in cells:
        H.set_cell_shading(c, bg)
        H.set_cell_borders(c)
    for ci, v in enumerate(r):
        H.add_run(cells[ci].paragraphs[0], v,
                  bold=(ci == 0 and accent),
                  size_pt=9.5,
                  color=(H.RED if ci == 0 and accent else H.CHARCOAL))


# ─── Page 9: What can be published now ─────────────────────────────────────
doc.add_page_break()
H.heading(doc, "7.  What can be published now", level=1)

H.para(doc, "Every publishable element is classified into one of four pathways. This is the steering view for government and MSG audiences.",
       size_pt=10.5, space_after=8)

H.centered_image(doc, CHARTS / "03-decision-buckets.png", width_cm=CONTENT_W - 1,
                 space_before=4, space_after=10)

buckets = [
    ("Ready to publish",
     "Fields the ZPPA e-GP system holds today, with clean source paths and clear legal authority.",
     "35 fields. Project IDs, titles, sectors, public authorities, procurement methods, tender and contract values, contractor IDs.",
     H.LOW_TEXT),
    ("Small format fixes",
     "Fields present but needing date-format conversion, codelist alignment, or free-text standardisation.",
     "20 fields. Free-text descriptions, sector-code harmonisation, ISO date conversion.",
     H.BLUE),
    ("Needs system work",
     "Fields NCC already collects under NCC Act section 53 that need an integration layer to reach the public.",
     "Approximately 70 NCC-held fields. Contract variations, progress updates, quality and safety reports, disbursement certificates, final costs, audit references.",
     H.YELLOW),
    ("Needs policy action",
     "Fields blocked by absence of a regulation, statutory instrument, or internal approval workflow.",
     "25 fields. Beneficial ownership; maintenance schedules; decommissioning plans; some audit and evaluation disclosures.",
     H.RED),
]

bt = doc.add_table(rows=len(buckets), cols=3)
H.set_table_full_width(bt, CONTENT_W)
H.apply_col_widths(bt, [4.2, 6.2, CONTENT_W - 10.4])
for ri, (label, what, example, col) in enumerate(buckets):
    cells = bt.rows[ri].cells
    H.set_row_height(bt.rows[ri], 2.0)
    H.set_cell_borders(cells[0], left="single", color=H.hex_of(col), sz="40")
    for c in cells:
        H.set_cell_shading(c, H.LGRAY)
        H.set_cell_margins(c, top=240, bottom=240, left=280, right=280)
    H.add_run(cells[0].paragraphs[0], label, bold=True, size_pt=11, color=col)
    H.add_run(cells[1].paragraphs[0], what, size_pt=10, color=H.CHARCOAL)
    H.add_run(cells[2].paragraphs[0], example, size_pt=10, color=H.CHARCOAL)

H.para(doc, "", space_after=10)


# ─── Page 11: What to do next ──────────────────────────────────────────────
doc.add_page_break()
H.heading(doc, "8.  What to do next", level=1)

H.para(doc, "Three disciplined moves follow directly from the mapping. Each is within reach inside a year. These are directional only. Detailed action items, named owners, target quarters, and budget ranges are out of scope for a field-level mapping report and should be developed separately through MSG consultation once the mapping is approved.",
       size_pt=10.5, space_after=10)

next_steps = [
    ("1.  Publish what is already publishable.",
     "The 35 fields in the ‘ready to publish’ bucket can go live through an OC4IDS feed on the existing e-GP portal. No new collection. No new law. The fastest observable progress this mapping can produce."),
    ("2.  Close the format-fix backlog.",
     "The 20 fields in the ‘small format fixes’ bucket need codelist alignment, date-format reconciliation, and free-text standardisation. Script-level work on the existing e-GP export."),
    ("3.  Open the NCC-to-ZPPA disclosure pathway.",
     "Most of the missing implementation data sits in NCC. A statutory instrument under NCC Act 2020 section 53 is the shortest path to integrate that data into the public feed. Scope needs country validation; legal drafting needs the Attorney-General’s Chambers."),
]
for title, body in next_steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    H.add_run(p, title + " ", bold=True, size_pt=10.5, color=H.CHARCOAL)
    H.add_run(p, body, size_pt=10.5, color=H.CHARCOAL)

H.para(doc, "", space_after=6)
H.callout_box(
    doc, CONTENT_W,
    label="FURTHER DETAIL",
    label_color=H.DMUTED,
    body="This report deliberately stops at three directional moves. A detailed reform programme (with legal crosswalk, named recommendations, proposed owners, target quarters, budget ranges, risks, and an implementation roadmap) is a separate product. Countries that want that strategic layer should commission it after MSG approval of this mapping report, not instead of it.",
)


# ─── Page 12: Review trail + conclusion ───────────────────────────────────
doc.add_page_break()
H.heading(doc, "9.  Review trail", level=1)
trail = [
    ("[Zambia Technical Team]", "Primary author (NCC and ZPPA, coordinated by CoST Zambia)",
     "[Date of lead author sign-off]"),
    ("CoST International Secretariat", "Field-level review against OC4IDS mapping methodology", "23 April 2026"),
    ("[CoST Zambia MSG]", "Multi-stakeholder review", "[scheduled]"),
    ("[ZPPA Director of IT]", "Technical feasibility sign-off on the e-GP data export and the OC4IDS feed", "[pending]"),
    ("[NCC Director of Construction Monitoring]", "Technical feasibility sign-off on NCC inspection-data integration and the source-paths table", "[pending]"),
]
trail_t = doc.add_table(rows=len(trail) + 1, cols=3)
H.set_table_full_width(trail_t, CONTENT_W)
H.apply_col_widths(trail_t, [6.0, 6.8, CONTENT_W - 12.8])
H.apply_default_padding(trail_t)
H.styled_table_header(trail_t.rows[0], ["Reviewer", "Role", "Date or status"])

for ri, r in enumerate(trail, start=1):
    cells = trail_t.rows[ri].cells
    H.set_row_height(trail_t.rows[ri], 1.2)
    bg = H.LGRAY if ri % 2 == 0 else H.WHITE
    for c in cells:
        H.set_cell_shading(c, bg)
        H.set_cell_borders(c)
    for ci, v in enumerate(r):
        H.add_run(cells[ci].paragraphs[0], v, size_pt=9.5,
                  color=H.CHARCOAL,
                  italic=(v.startswith("[") and v.endswith("]")))


# ─── Page 13: Conclusion + annexes ─────────────────────────────────────────
doc.add_page_break()
H.heading(doc, "10.  Conclusion", level=1)

H.para(doc, "Based on this mapping, Zambia's public infrastructure data system is not as thin as it first appears. The ZPPA e-GP holds the procurement backbone. The NCC holds the implementation record. The Access to Information Act 2023 appears to mandate publication of most of both. The binding constraint, on the evidence reviewed, is integration and disclosure rather than collection.",
       size_pt=10.5, space_after=8)
H.para(doc, "This analysis identifies 35 OC4IDS fields that can move to public disclosure in the current quarter. Another 20 would follow after small format fixes. A further group of NCC-held implementation fields (approximately 70) would become publishable through an integration layer. The remainder require regulation or workflow change.",
       size_pt=10.5, space_after=8)
H.para(doc, "The immediate path forward has three moves: publish what is already publishable, draft and sign an NCC and ZPPA data-sharing instrument, and scope the integration work as a publication pathway rather than a monitoring system. Scoping, sequencing, and institutional design sit outside this mapping report and should be picked up in a separate planning conversation.",
       size_pt=10.5, bold=True, color=H.CHARCOAL, space_after=14)

H.heading(doc, "Annexes", level=2)

# Scope note sits above the bullet list, not inside it, so it is clearly not
# itself an annex.
H.para(doc, "Note on scope: legal crosswalk, detailed recommendations, implementation roadmap, risks register, and NCC inspection form redesign are deliberately out of scope for a field-level mapping report. Countries that want that layer should commission it as a separate reform-planning product.",
       size_pt=9.5, italic=True, color=H.DMUTED, space_after=8)

for a in [
    "Annex A.  Completed OC4IDS v0.9.5 field-level mapping template (attached: 2025_11_Zambia_OC4IDS_0.9.5_Field-Level_Mapping_Template.xlsx).",
    "Annex B.  Full source-paths table: every mapped OC4IDS path crossed with its ZPPA e-GP source path.",
    "Annex C.  Country-specific fields preserved as extensions.",
    "Annex D.  Sampled procurement records used for field-population verification (n = 50).",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    H.add_run(p, "•  ", color=H.MUTED, size_pt=10.5)
    H.add_run(p, a, size_pt=10, color=H.CHARCOAL)

H.para(doc, "", space_after=14)


# Placeholder checklist is NOT part of the published sample. It is scaffolding
# that belongs in an internal pre-publication tool, not in the model document
# a country team will reference. Written as a side-car Markdown file next to
# the DOCX; keep it, delete it, or email it to the MSG as a checklist.
_placeholder_items = [
    ("Cover block", "[Lead author: to be confirmed by MSG]",
     "Publisher / Prepared-by block on the cover page."),
    ("Cover block", "[portal URL to be set at IPPI-Zambia launch]",
     "Publication-licence row on the cover page."),
    ("Cover block", "[to be set by Zambia team]",
     "Version line on the cover page (date)."),
    ("Section 1.2", "[Note to Zambia team] on sampling",
     "Opening paragraph note on record count, sector mix, value totals. Replace with actual universe figures or delete if universe review holds."),
    ("Section 9 (Review trail)", "[Zambia Technical Team] and [Date of lead author sign-off]",
     "Replace with the actual lead author name and sign-off date."),
    ("Section 9 (Review trail)", "[CoST Zambia MSG], status [scheduled]",
     "Set the MSG review date once scheduled."),
    ("Section 9 (Review trail)", "[ZPPA Director of IT], status [pending]",
     "Replace with the actual officer and confirm sign-off."),
    ("Section 9 (Review trail)", "[NCC Director of Construction Monitoring], status [pending]",
     "Replace with the actual officer and confirm sign-off."),
]

_checklist_path = OUT.parent / "02-sample-final-report_placeholder-checklist.md"
_lines = [
    "# Pre-publication placeholder checklist",
    "",
    "Internal tool. Not part of the published sample.",
    "",
    "Every bracketed placeholder in 02-sample-final-report.docx. Tick through each row before the MSG signs the report for publication. A single un-replaced placeholder in a public report damages its standing; a cheap tick-through on draft day is cheaper than republishing a corrected version after distribution.",
    "",
    "| Location | Placeholder | Action required |",
    "|----------|-------------|-----------------|",
]
for loc, tag, action in _placeholder_items:
    _lines.append(f"| {loc} | `{tag}` | {action} |")
_checklist_path.write_text("\n".join(_lines) + "\n", encoding="utf-8")

H.para(doc, "CoST, the Infrastructure Transparency Initiative  •  www.infrastructuretransparency.org  •  Published under CC BY 4.0",
       size_pt=8, italic=True, color=H.LMUTED,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)


doc.save(str(OUT))
print(f"OK {OUT}")
