"""Build Doc 1: CoST-branded review letter for Zambia OC4IDS report.

Archetype: Review/Feedback. Intensity: Professional Report. Layout: Landscape A4.
"""
import pathlib
from docx import Document
from docx.shared import Cm, Mm, Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ─── CoST palette ───────────────────────────────────────────────────────────
RED = RGBColor(0xB7, 0x25, 0x1C)
COVER_RED = RGBColor(0xCC, 0x20, 0x28)
CHARCOAL = RGBColor(0x4A, 0x47, 0x43)
BLUE = RGBColor(0x71, 0xB6, 0xC9)
YELLOW = RGBColor(0xF1, 0xC0, 0x3D)
MUTED = RGBColor(0x9C, 0x9B, 0x99)
DMUTED = RGBColor(0x6B, 0x6A, 0x68)
LMUTED = RGBColor(0xB0, 0xB0, 0xB1)
MGRAY = RGBColor(0xE5, 0xE4, 0xE3)
LGRAY = RGBColor(0xF3, 0xF2, 0xF1)
LIGHT_NEUTRAL = RGBColor(0xF0, 0xED, 0xE7)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Priority tints
CRITICAL_BG = RGBColor(0xF5, 0xD5, 0xD3)
HIGH_BG = RGBColor(0xD9, 0xED, 0xF3)
MEDIUM_BG = RGBColor(0xFD, 0xF3, 0xD7)
LOW_BG = RGBColor(0xE8, 0xF5, 0xE9)
MEDIUM_TEXT = RGBColor(0xD4, 0xA0, 0x17)
LOW_TEXT = RGBColor(0x4E, 0xA8, 0x3D)

ASSETS = pathlib.Path("/Users/cengkurumichael/.claude/skills/cost-document-design/assets")
CHARTS = pathlib.Path(__file__).parent.parent / "charts"
OUT = pathlib.Path(__file__).parent.parent / "01-review-letter.docx"

# ─── docx helpers ───────────────────────────────────────────────────────────

def hex_of(c):
    """Convert RGBColor or 6-char string to 6-char hex."""
    if isinstance(c, str):
        return c.upper().lstrip("#")
    return "".join(f"{b:02X}" for b in c)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_of(hex_color))
    tc_pr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, color="E5E4E3", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), val if val else 'nil')
        if val:
            border.set(qn('w:sz'), sz)
            border.set(qn('w:color'), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_run(p, text, *, bold=False, italic=False, size_pt=10.5, color=CHARCOAL, font="Arial"):
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def para(doc, text, *, style=None, size_pt=10.5, color=CHARCOAL, bold=False, italic=False,
         space_after=6, align=None, font="Arial"):
    p = doc.add_paragraph() if style is None else doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if text:
        add_run(p, text, bold=bold, italic=italic, size_pt=size_pt, color=color, font=font)
    return p


def accent_bar(doc, color=RED, width_pt=100):
    """A thin coloured rule, implemented as a 1-row, 1-col table with bottom border."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.allow_autofit = False
    cell = t.cell(0, 0)
    cell.text = ""
    # width via column
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_pt * 20))
    tblW.set(qn('w:type'), 'dxa')
    t._tbl.tblPr.append(tblW)
    set_cell_borders(cell, bottom="single",
                     color=hex_of(color),
                     sz="24")
    return t


def set_landscape(section, margin_mm=18):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(margin_mm)
    section.bottom_margin = Mm(margin_mm)
    section.left_margin = Mm(margin_mm)
    section.right_margin = Mm(margin_mm)


# ─── Document ───────────────────────────────────────────────────────────────

doc = Document()

# Page size + margins
section = doc.sections[0]
set_landscape(section)

# Default style
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)
style.font.color.rgb = CHARCOAL

# --- Header: CoST logo right-aligned ---
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hp.paragraph_format.space_after = Pt(0)
hr = hp.add_run()
hr.add_picture(str(ASSETS / "cost-logo-real.jpeg"), width=Cm(4.2))

# --- Footer: stripe + page number ---
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
fp.paragraph_format.space_before = Pt(0)
fp.paragraph_format.space_after = Pt(0)
fr = fp.add_run()
fr.add_picture(str(ASSETS / "cost-stripe-real.png"), width=Cm(16), height=Mm(3.5))
fp2 = footer.add_paragraph()
fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fp2.paragraph_format.space_before = Pt(4)
fld_run = fp2.add_run()
fld_run.font.name = 'Arial'
fld_run.font.size = Pt(8.5)
fld_run.font.color.rgb = MUTED
# PAGE field
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'PAGE')
fld_run._element.append(fld)

# ─── Page 1: Title block ───────────────────────────────────────────────────
para(doc, "OVERLAY REVIEW", size_pt=11, bold=True, color=RED, space_after=2)
para(doc, "OC4IDS Field-Level Mapping Report", size_pt=22, bold=True, color=CHARCOAL, space_after=4)
para(doc, "Zambia (ZPPA / NCC / CoST Zambia)  ·  Draft of 3 March 2026", size_pt=13, color=DMUTED, space_after=12)
accent_bar(doc, RED, width_pt=380)
para(doc, "", space_after=6)

# Metadata block
meta_t = doc.add_table(rows=4, cols=2)
meta_t.autofit = False
meta_t.allow_autofit = False
for row in meta_t.rows:
    row.cells[0].width = Cm(4.5)
    row.cells[1].width = Cm(18)
meta_rows = [
    ("Review framework:", "Overlay v0.9  ·  Reviewer Gate Check (Part B) + Full Rubric (Part C)"),
    ("Reviewed by:", "Michael Cengkuru  ·  CoST International Secretariat"),
    ("Review date:", "2026-04-23"),
    ("Decision:", "REVISE AND RESUBMIT  ·  Gate score 3 / 10"),
]
for i, (k, v) in enumerate(meta_rows):
    kc, vc = meta_t.rows[i].cells
    kp = kc.paragraphs[0]
    add_run(kp, k, bold=True, size_pt=10, color=DMUTED)
    vp = vc.paragraphs[0]
    if i == 3:
        add_run(vp, "REVISE AND RESUBMIT", bold=True, size_pt=10.5, color=RED)
        add_run(vp, "  ·  Gate score 3 / 10", size_pt=10.5, color=CHARCOAL)
    else:
        add_run(vp, v, size_pt=10.5, color=CHARCOAL)
    for c in (kc, vc):
        set_cell_borders(c)
para(doc, "", space_after=12)

# ─── Verdict box ───────────────────────────────────────────────────────────
vb = doc.add_table(rows=1, cols=1)
vb.autofit = False
vc = vb.cell(0, 0)
set_cell_shading(vc, LGRAY)
set_cell_borders(vc, top="single", bottom="single",
                 color=hex_of(RED), sz="24")
vcp = vc.paragraphs[0]
add_run(vcp, "VERDICT   ", bold=True, size_pt=10, color=RED)
add_run(vcp, "Zambia has produced a materially stronger mapping template than the Kaduna calibration floor ",
        size_pt=10.5, color=CHARCOAL)
add_run(vcp, "(73 mapped fields, all substantive, 97.9% source forward-mapping). ",
        bold=True, size_pt=10.5, color=CHARCOAL)
add_run(vcp, "The narrative report does not surface any of these numbers. The legal framework analysis in Annex 8.1 is the strongest we have seen on this kind of report. Fix the quantitative layer and this report reaches a solid band.",
        size_pt=10.5, color=CHARCOAL)
vcp2 = vc.add_paragraph()
vcp2.paragraph_format.space_before = Pt(4)
add_run(vcp2, "Resubmission target: 2026-05-21.", bold=True, size_pt=10, color=CHARCOAL)
para(doc, "", space_after=14)

# ─── Headline findings chart ───────────────────────────────────────────────
sec = doc.add_paragraph()
add_run(sec, "1.  The template evidence", bold=True, size_pt=13, color=CHARCOAL)
accent_bar(doc, RED, width_pt=160)
para(doc, "", space_after=4)

para(doc, "Zambia mapped 73 of 1,480 OC4IDS template slots across four sheets. All mapped fields are substantive (no placeholders). The Projects and Contracting Processes sheets carry the bulk of the coverage; Linked Releases is at zero and Parties is effectively silent.",
     size_pt=10.5, space_after=8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run()
r.add_picture(str(CHARTS / "01-sheet-coverage.png"), width=Cm(18))
para(doc, "", space_after=6)

# ─── Gate scorecard ────────────────────────────────────────────────────────
sec = doc.add_paragraph()
add_run(sec, "2.  Gate scorecard", bold=True, size_pt=13, color=CHARCOAL)
accent_bar(doc, RED, width_pt=160)
para(doc, "", space_after=4)

para(doc, "Overlay's 10-item gate check is pass/fail. A score below 7 returns the report at triage without full-rubric scoring. Zambia scored 3 / 10.",
     size_pt=10.5, space_after=8)

gate_rows = [
    ("1", "Quantified headline (%)", "FAIL", "Executive summary states \"significant data gaps\". No percentage anywhere."),
    ("2", "Portal URLs + access dates", "FAIL", "ZPPA e-GP named; URL (eprocure.zppa.org.zm) lives only in the template. No access date."),
    ("3", "Sample defined", "FAIL", "No project count, value, sectors, time window, or sampling method."),
    ("4", "OC4IDS template version", "PASS", "\"OC4IDS v0.9.5\" stated in §3 Methodology."),
    ("5", "Coverage % by phase", "FAIL", "Phases discussed qualitatively. Template supports percentages directly."),
    ("6", "Source-field provenance", "FAIL", "Template has clean notation (zppa_egp.project.project.id). Report buries it."),
    ("7", "Gap root causes (typology)", "FAIL", "Gaps labelled \"missing\". The NCC-already-collects-this insight is a type-6 gap (collected, not disclosed) and should be labelled."),
    ("8", "Procurement + FOI law cited", "PASS", "Strong: ATI Act s.6/8/9/17, ZPPA Act s.67/70, NCC Act s.5/31/33/53."),
    ("9", "Recommendations with owners + timelines", "FAIL", "Six groups, all directional. No owner, no timeline, no effort level."),
    ("10", "Template attached", "PASS", "Workbook attached alongside."),
]

t = doc.add_table(rows=len(gate_rows) + 1, cols=4)
t.autofit = False
widths = [Cm(1.0), Cm(5.8), Cm(2.0), Cm(15.8)]
for row in t.rows:
    for c, w in zip(row.cells, widths):
        c.width = w

# Header
for i, h in enumerate(["#", "Gate item", "Result", "Note"]):
    c = t.rows[0].cells[i]
    set_cell_shading(c, "2D2D2D")
    set_cell_borders(c, top="single", bottom="single", left="single", right="single", color="2D2D2D", sz="4")
    hp = c.paragraphs[0]
    add_run(hp, h, bold=True, size_pt=10, color=WHITE)

# Body
for ri, (num, item, result, note) in enumerate(gate_rows, start=1):
    row_cells = t.rows[ri].cells
    bg = LGRAY if ri % 2 == 0 else WHITE
    for c in row_cells:
        set_cell_shading(c, bg)
        set_cell_borders(c)
    add_run(row_cells[0].paragraphs[0], num, size_pt=10, color=MUTED)
    add_run(row_cells[1].paragraphs[0], item, size_pt=10, color=CHARCOAL)
    # result cell: colour-coded
    rp = row_cells[2].paragraphs[0]
    if result == "PASS":
        add_run(rp, result, bold=True, size_pt=10, color=LOW_TEXT)
        set_cell_shading(row_cells[2], LOW_BG)
    else:
        add_run(rp, result, bold=True, size_pt=10, color=RED)
        set_cell_shading(row_cells[2], CRITICAL_BG)
    add_run(row_cells[3].paragraphs[0], note, size_pt=9.5, color=CHARCOAL)

para(doc, "", space_after=4)
para(doc, "Spot-sample (template evidence check): 5 of 5 mapped rows carry substantive content, no placeholders. This is the strongest spot-sample result we have seen. Unmapped rows, however, have empty Notes columns: gate rule fails on unjustified omissions.",
     size_pt=10, italic=True, color=DMUTED, space_after=12)

# ─── Phase coverage chart ──────────────────────────────────────────────────
doc.add_page_break()
sec = doc.add_paragraph()
add_run(sec, "3.  Where the gaps sit", bold=True, size_pt=13, color=CHARCOAL)
accent_bar(doc, RED, width_pt=160)
para(doc, "", space_after=4)

para(doc, "Coverage is front-loaded on identification and procurement. The middle of the lifecycle, where value-for-money judgments live, is almost empty. Maintenance and decommissioning return zero mapped fields.",
     size_pt=10.5, space_after=8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(CHARTS / "02-phase-coverage.png"), width=Cm(18))
para(doc, "", space_after=6)

para(doc, "The report's §5.3 labels Implementation a \"critical gap\" and §4.1 flags that the NCC already collects most of the missing data through its monitoring mandate (NCC Act 2020, s.53). Under the Overlay typology this is a type-6 gap (collected, not disclosed). Remediation is data integration, not data collection. Label it.",
     size_pt=10.5, space_after=14)

# ─── Findings table ────────────────────────────────────────────────────────
sec = doc.add_paragraph()
add_run(sec, "4.  Findings, owners, and specific revisions", bold=True, size_pt=13, color=CHARCOAL)
accent_bar(doc, RED, width_pt=160)
para(doc, "", space_after=4)

findings = [
    ("F1", "Critical", "Executive summary carries no quantified headline.",
     "\"Significant data gaps exist\" is the current phrasing. The template supports \"73 of 1,480 OC4IDS fields mapped (4.9%); zero coverage on Linked Releases and effectively zero on Parties\".",
     "Rewrite the opening sentence with the percentage and the four per-sheet counts. Match Part A item 1.",
     "Zambia Technical Team", "2026-05-07"),
    ("F2", "Critical", "Phase-level percentages not disclosed.",
     "Phases are discussed qualitatively. The template allows quantification: Identification 28%, Procurement 21%, Preparation 12%, Completion 8%, Implementation 2%, Maintenance 0%, Decommissioning 0%.",
     "Add a phase coverage table to §4 Findings with one row per phase, a percentage column, and the dominant gap type.",
     "Zambia Technical Team", "2026-05-10"),
    ("F3", "Critical", "Source-field provenance buried in the template.",
     "Every mapped OC4IDS path in the template has a clean source notation, e.g. `/id` -> `zppa_egp.project.project.id`. The report does not expose these mappings.",
     "Add a provenance column to the §5 phase tables: OC4IDS path, ZPPA e-GP path, status.",
     "Zambia Technical Team", "2026-05-14"),
    ("F4", "Critical", "Recommendations have no owners or timelines.",
     "§6.1 to §6.6 list six recommendation groups. All are directional. None name an institution or a quarter. Overlay gate item 9 fails.",
     "Assign each recommendation to NCC, ZPPA, CoST Zambia, or MoF with a specific quarter. The NCC Act and ATI Act already provide legal mandates; cite them per recommendation.",
     "CoST Zambia + NCC + ZPPA", "2026-05-14"),
    ("F5", "High", "Gap typology not applied.",
     "Gaps are described by phase, not by root cause. The strongest implicit finding - \"NCC already collects this data\" - is a type-6 gap (collected, not disclosed). Remediation differs by type.",
     "Apply the 8-category typology to each material gap in §5. Label the NCC insight explicitly as type-6; this reframes §6 from \"collect more data\" to \"integrate and disclose\".",
     "Zambia Technical Team", "2026-05-10"),
    ("F6", "High", "Data beyond OC4IDS not examined.",
     "The report does not discuss which ZPPA e-GP or NCC fields fall outside OC4IDS. Country-specific disclosure that could feed upstream to OCP is invisible. Scored in Overlay C9.",
     "Add a short section on non-template fields: what is disclosed locally, whether to preserve as extensions, what might be proposed for OC4IDS standard revision.",
     "Zambia Technical Team + CoST IS", "2026-05-21"),
    ("F7", "High", "Decision summary panel absent.",
     "Annex 8.2 is a first pass at the lifecycle-to-data-source crosswalk, but it does not classify elements as publish-now / light-transform / system-change / legal-action. C11's mandatory decision panel is missing.",
     "Add a decision summary panel to §7 Conclusion using the four-bucket classification. See the sample corrected report, §8, for the target format.",
     "Zambia Technical Team", "2026-05-17"),
    ("F8", "Medium", "Portal URL and access dates missing from the body.",
     "The ZPPA e-GP URL (eprocure.zppa.org.zm/epps/home.d) appears only in the template's (Source) 1 sheet.",
     "Add the URL and the date(s) accessed to §3.1 Data Sources.",
     "Zambia Technical Team", "2026-05-07"),
    ("F9", "Medium", "Sample not quantified.",
     "\"The ZPPA e-GP system was analysed\" does not specify the scope.",
     "State the count of procurements reviewed, aggregate value, sectors covered, and the time window.",
     "Zambia Technical Team", "2026-05-07"),
    ("F10", "Low", "Unjustified omissions in template Notes column.",
     "Spot-sample of 5 unmapped template rows found no justification in the Notes column. Overlay rule fails on silent omissions.",
     "Add a Notes entry to every unmapped row in the (OC4IDS) sheets before resubmission. A one-line reason is enough.",
     "Zambia Technical Team", "2026-05-14"),
]

ft = doc.add_table(rows=len(findings) + 1, cols=6)
ft.autofit = False
col_w = [Cm(0.9), Cm(1.8), Cm(5.0), Cm(7.8), Cm(6.3), Cm(3.4)]
for row in ft.rows:
    for c, w in zip(row.cells, col_w):
        c.width = w
headers = ["Ref", "Priority", "Issue", "Evidence", "Recommendation", "Owner / deadline"]
for i, h in enumerate(headers):
    c = ft.rows[0].cells[i]
    set_cell_shading(c, "2D2D2D")
    set_cell_borders(c, top="single", bottom="single", left="single", right="single", color="2D2D2D")
    add_run(c.paragraphs[0], h, bold=True, size_pt=10, color=WHITE)

prio_map = {
    "Critical": (RED, CRITICAL_BG),
    "High": (BLUE, HIGH_BG),
    "Medium": (MEDIUM_TEXT, MEDIUM_BG),
    "Low": (LOW_TEXT, LOW_BG),
}
for ri, (ref, prio, issue, ev, rec, owner, dl) in enumerate(findings, start=1):
    cells = ft.rows[ri].cells
    bg = LGRAY if ri % 2 == 0 else WHITE
    for c in cells:
        set_cell_shading(c, bg)
        set_cell_borders(c)
    add_run(cells[0].paragraphs[0], ref, bold=True, size_pt=9.5, color=CHARCOAL)
    pc_text_color, pc_bg = prio_map[prio]
    set_cell_shading(cells[1], pc_bg)
    add_run(cells[1].paragraphs[0], prio, bold=True, size_pt=9.5, color=pc_text_color)
    add_run(cells[2].paragraphs[0], issue, bold=True, size_pt=9.5, color=CHARCOAL)
    add_run(cells[3].paragraphs[0], ev, size_pt=9.5, color=CHARCOAL)
    add_run(cells[4].paragraphs[0], rec, size_pt=9.5, color=CHARCOAL)
    owner_p = cells[5].paragraphs[0]
    add_run(owner_p, owner, size_pt=9.5, color=CHARCOAL)
    p2 = cells[5].add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    add_run(p2, dl, bold=True, size_pt=9.5, color=RED)

para(doc, "", space_after=12)

# ─── Decision buckets chart ────────────────────────────────────────────────
doc.add_page_break()
sec = doc.add_paragraph()
add_run(sec, "5.  What Zambia can publish now", bold=True, size_pt=13, color=CHARCOAL)
accent_bar(doc, RED, width_pt=160)
para(doc, "", space_after=4)

para(doc, "The four-bucket classification below is the steering view government and MSG readers actually use. It sits in §C11 of the Overlay rubric as a mandatory panel. The report's §8.2 Lifecycle Framework is the foundation; expand it into this classification.",
     size_pt=10.5, space_after=8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(CHARTS / "03-decision-buckets.png"), width=Cm(18))
para(doc, "", space_after=6)

# ─── Legal crosswalk chart ─────────────────────────────────────────────────
sec = doc.add_paragraph()
add_run(sec, "6.  The legal lever", bold=True, size_pt=13, color=CHARCOAL)
accent_bar(doc, RED, width_pt=160)
para(doc, "", space_after=4)

para(doc, "Zambia's legal framework already authorises publication of most OC4IDS fields. Annex 8.1 is publication-quality and should anchor the executive summary, not sit at the back of the report.",
     size_pt=10.5, space_after=8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(CHARTS / "04-legal-crosswalk.png"), width=Cm(18))
para(doc, "", space_after=14)

# ─── Strengths to preserve ─────────────────────────────────────────────────
sec = doc.add_paragraph()
add_run(sec, "7.  Strengths to preserve on revision", bold=True, size_pt=13, color=CHARCOAL)
accent_bar(doc, LOW_TEXT, width_pt=160)
para(doc, "", space_after=6)

strengths = [
    ("Legal framework analysis (Annex 8.1).",
     "ATI Act 2023 s.6/8/9/17 with quoted provisions, ZPPA Act No.3/2020 s.67/70, NCC Act No.10/2020 s.5/31/33/53 with subsections. The strongest legal anchoring we have seen on this review cycle. Move the essence of this analysis into the executive summary, keep the full text in the annex."),
    ("The NCC insight.",
     "The observation that most implementation-stage data is already collected by NCC under its statutory monitoring mandate is the single most actionable finding in the report. Label it explicitly as a type-6 gap (collected, not disclosed) and anchor §6 recommendations around integration rather than collection."),
    ("Source forward-mapping.",
     "The template maps 47 of 48 source data elements forward to OC4IDS (97.9%). This is exceptional. Lift the number directly into §3 Methodology."),
    ("Lifecycle framework (§8.2).",
     "The five-row data-source-by-stage crosswalk is the skeleton of the decision summary panel in §C11 of the rubric. Expand it into the four-bucket classification (publish-now / light-transform / system-change / legal-action)."),
    ("No placeholder mappings.",
     "Spot-sample of 5 mapped rows returned 5 substantive entries. This is the strongest evidence-check result on record for Overlay. Do not change the template mappings during revision, the numbers are correct."),
]
for title, body in strengths:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "+  ", bold=True, size_pt=10.5, color=LOW_TEXT)
    add_run(p, title + "  ", bold=True, size_pt=10.5, color=CHARCOAL)
    add_run(p, body, size_pt=10.5, color=CHARCOAL)

para(doc, "", space_after=14)

# ─── Companion document ────────────────────────────────────────────────────
sec = doc.add_paragraph()
add_run(sec, "8.  Companion document", bold=True, size_pt=13, color=CHARCOAL)
accent_bar(doc, RED, width_pt=160)
para(doc, "", space_after=6)

cb = doc.add_table(rows=1, cols=1)
cc = cb.cell(0, 0)
set_cell_shading(cc, LGRAY)
set_cell_borders(cc, left="single", color=hex_of(BLUE), sz="32")
ccp = cc.paragraphs[0]
add_run(ccp, "02-sample-corrected-report.docx", bold=True, size_pt=10.5, color=BLUE)
ccp2 = cc.add_paragraph()
add_run(ccp2, "A full rewrite of the Zambia report showing the target state after all revisions above are applied. Use it as a reference; do not copy it wholesale. Preserve your legal analysis, the NCC insight, and the lifecycle framework; rewrite the executive summary, findings tables, and recommendations sections against the corrected structure.",
        size_pt=10, color=CHARCOAL)
para(doc, "", space_after=12)

# ─── Footer signature ──────────────────────────────────────────────────────
para(doc, "Reference: Overlay v0.9  ·  Reviewer Gate Check (Part B v1.0) + Full Rubric (Part C v0.9)",
     size_pt=9, italic=True, color=MUTED, space_after=0)
para(doc, "Sources: Zambia OC4IDS Field-Level Mapping Template v0.9.5, accessed 2026-04-23. n=1,480 template slots. Zambia OC4IDS Field-Level Mapping Report (draft), 3 March 2026.",
     size_pt=9, italic=True, color=MUTED, space_after=0)

doc.save(str(OUT))
print(f"OK {OUT}")
