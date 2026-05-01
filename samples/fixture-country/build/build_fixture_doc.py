"""Minimal build script for the fixture-country test package.

Writes a DOCX that intentionally contains some mutable weaknesses so
the auto-researcher has something to propose against. The content is
deliberately synthetic — no real country data.
"""
from __future__ import annotations
import pathlib

from docx import Document


HERE = pathlib.Path(__file__).resolve().parent
OUT = HERE.parent / "fixture-report.docx"


def build():
    d = Document()

    # Headline section — mentions "100" as the denominator, which per
    # fixture-country's manifest is the forbidden literal, but this
    # paragraph is under "Findings" (neutral) and a permitted callout.
    d.add_heading("Findings", level=1)
    d.add_paragraph(
        "This review covers 25 of 100 mapped fields. "
        "Recommendations R1 and R2 reference the action list."
    )
    d.add_paragraph("NOTE: the total of 100 fields is the universe.")

    # Phase table text
    d.add_heading("Phases", level=1)
    d.add_paragraph("Alpha: 50%. Beta: 25%.")

    # A deliberate AI-tell and internal-vocab term so the auto-loop
    # can propose mutations. No dashes used.
    d.add_paragraph(
        "We will leverage this review to examine integration paths."
    )

    d.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
